#!/usr/bin/env python3
"""
DMAP Medium Article Generation Script (English Version)
Generates a Medium article in Word (.docx) format using python-docx.

Usage: python docs/article/generate_medium_article_en.py
Output: docs/article/DMAP-medium-article-en.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images", "en")
IMAGES_DIR_FALLBACK = os.path.join(SCRIPT_DIR, "images")  # Korean images as fallback
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "DMAP-medium-article-en.docx")

# ---------------------------------------------------------------------------
# Style Constants
# ---------------------------------------------------------------------------
FONT_NAME = "Arial"
FONT_FALLBACK = "Arial"

# Colors
COLOR_HEADING_H1 = "1A1A1A"
COLOR_HEADING_H2 = "2D2D2D"
COLOR_SUBTITLE = "555555"
COLOR_TABLE_HEADER_BG = "D5E8F0"
COLOR_TABLE_BORDER = "CCCCCC"
COLOR_NOTE_BG = "F0F4F8"
COLOR_NOTE_BORDER = "4A6FA5"
COLOR_ACCENT = "2C5282"


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add table borders (color: #CCCCCC)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}></w:tblPr>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table (blue header background #D5E8F0)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = FONT_NAME

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_image_if_exists(doc, filename, caption, width_inches=5.5):
    """Insert an image (falls back to Korean images if English version missing)."""
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        # Fallback to Korean images directory
        path = os.path.join(IMAGES_DIR_FALLBACK, filename)
        if not os.path.exists(path):
            print(f"  [WARNING] Image missing: {filename}")
            p = doc.add_paragraph()
            run = p.add_run(f"[Missing image: {filename}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        print(f"  [INFO] Using Korean fallback image: {filename}")
    else:
        print(f"  [OK] English image found: {filename}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run.italic = True


def add_heading(doc, text, level=1):
    """Add a heading (H1=16pt color #1A1A1A, H2=13pt color #2D2D2D)."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(
                int(COLOR_HEADING_H1[0:2], 16),
                int(COLOR_HEADING_H1[2:4], 16),
                int(COLOR_HEADING_H1[4:6], 16)
            )
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(
                int(COLOR_HEADING_H2[0:2], 16),
                int(COLOR_HEADING_H2[2:4], 16),
                int(COLOR_HEADING_H2[4:6], 16)
            )
    return h


def add_body(doc, text):
    """Add body text (Arial, 11pt, 1.5 line spacing)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


def add_body_bold(doc, text):
    """Add bold body text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    return p


def add_note(doc, text, label="NOTE"):
    """NOTE box (1x1 table, background #F0F4F8, left thick border #4A6FA5 24pt)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(
        int(COLOR_NOTE_BORDER[0:2], 16),
        int(COLOR_NOTE_BORDER[2:4], 16),
        int(COLOR_NOTE_BORDER[4:6], 16)
    )
    run.font.name = FONT_NAME
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    set_cell_shading(cell, COLOR_NOTE_BG)
    # Left thick border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:left w:val="single" w:sz="24" w:color="{COLOR_NOTE_BORDER}"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)


def add_bullet(doc, text):
    """Add a bullet list item."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


# ---------------------------------------------------------------------------
# Section functions
# ---------------------------------------------------------------------------

def section_01_intro(doc):
    """Section 1: Introduction."""
    add_heading(doc, "The Structural Challenge of Multi-Agent Systems", level=1)

    add_body(
        doc,
        "When LLM agents expand from single entities to collaborative multi-agent systems, "
        "complexity increases not linearly but combinatorially. The number of relationships "
        "to manage \u2014 role boundaries between agents, tool access permissions, execution "
        "order, escalation conditions \u2014 grows proportionally to the square of the number "
        "of agents, creating a fundamental structural problem."
    )

    add_body(
        doc,
        "Existing approaches such as LangChain, CrewAI, and AutoGen address this problem "
        "through Python/TypeScript SDK code. Agent definitions, tool connections, and "
        "orchestration logic are all intertwined within programming languages, forming a "
        "structure tightly coupled to specific SDKs. As a result, changing the runtime "
        "environment requires rewriting the entire system, and domain experts or "
        "non-developers are excluded from participating in agent system design."
    )

    add_body(
        doc,
        "This raises a fundamental question: is code truly necessary to define agent roles "
        "and relationships? Declaring \"what an agent can do\" and coding \"how it operates\" "
        "are inherently different activities."
    )

    doc.add_paragraph()


def section_02_declarative(doc):
    """Section 2: The Declarative Approach."""
    add_heading(doc, "The Essence of the Declarative Approach \u2014 \"What\" vs. \"How\"", level=1)

    add_body(
        doc,
        "Declarative programming is a paradigm that describes \"the desired outcome\" and "
        "delegates \"how to achieve it\" to the execution environment. SQL declares data "
        "conditions without specifying search algorithms; HTML/CSS declares document "
        "structure and style without dictating rendering engine implementation; Kubernetes "
        "YAML declares desired infrastructure state without prescribing container scheduling "
        "methods. Their common thread lies in the separation of execution environment "
        "(runtime) from declaration (specification)."
    )

    add_body(
        doc,
        "The same principle can be applied to multi-agent systems. Agent roles, capabilities, "
        "and constraints are declared in Markdown and YAML, while execution is delegated to "
        "runtimes (Claude Code, Codex CLI, Gemini CLI, etc.). The core value of this approach "
        "can be summarized in two points. First, runtime neutrality \u2014 an abstract "
        "declaration like `tier: HIGH` can be interpreted as the optimal model in any runtime "
        "environment. Second, non-developer accessibility \u2014 anyone who knows Markdown "
        "and YAML can build agent systems, including domain experts."
    )

    make_table(
        doc,
        headers=["Comparison", "Traditional Frameworks", "Declarative Approach"],
        rows=[
            ["Agent Definition", "Python/TypeScript SDK code", "Markdown + YAML"],
            ["Orchestration", "Graph/function call chain code", "Skill prompts (natural language)"],
            ["Runtime Dependency", "Tightly coupled to specific SDK", "Runtime-neutral"],
            ["Architecture Principles", "None or framework-dependent", "Clean Architecture"],
            ["Tool Integration", "Tool object creation in code", "Abstract declaration (tools.yaml) \u2192 Gateway mapping"],
            ["Tier Management", "None", "4-Tier (HEAVY/HIGH/MEDIUM/LOW) + automatic runtime mapping"],
            ["Portability", "Low (full rewrite required)", "High (only replace runtime-mapping.yaml)"],
        ],
        col_widths=[4, 6, 6],
    )

    doc.add_paragraph()


def section_03_architecture(doc):
    """Section 3: Architecture."""
    add_heading(doc, "Architecture Design \u2014 Applying Clean Architecture to Agents", level=1)

    add_body(
        doc,
        "The core principles of Clean Architecture \u2014 unidirectional dependency and "
        "separation of concerns \u2014 are proven design philosophies that ensure long-term "
        "maintainability of software systems. The declarative multi-agent architecture applies "
        "these principles directly to agent systems, organized into three core layers."
    )

    add_body_bold(doc, "Skills (Controller Layer)")
    add_body(
        doc,
        "The entry point for user requests and the orchestration layer. It classifies request "
        "intent and performs routing by delegating tasks to appropriate agents. Delegation "
        "skills (Core, Planning, Orchestrator) invoke agents, while direct skills (Setup, "
        "Utility) use the Gateway directly. The YAGNI principle is applied by skipping "
        "unnecessary layers depending on the execution path."
    )

    add_body_bold(doc, "Agents (Service Layer)")
    add_body(
        doc,
        "Role-based autonomous execution units. Each agent consists of a pair: AGENT.md "
        "(goals and workflow prompts) and agentcard.yaml (identity, capabilities, and "
        "constraint declarations). Each agent focuses on a single specialized role, and "
        "tasks that cross role boundaries are delegated to other agents according to "
        "handoff rules."
    )

    add_body_bold(doc, "Gateway (Infrastructure Layer)")
    add_body(
        doc,
        "The infrastructure layer that provides mapping tables between abstract declarations "
        "and concrete runtimes. runtime-mapping.yaml transforms agent abstract tiers into "
        "actual LLM models and abstract tools into actual tool implementations. When the "
        "runtime environment changes, only this mapping file needs to be replaced."
    )

    add_body(
        doc,
        "Beyond these three layers, a Resource Marketplace serves as a shared pool of guides, "
        "templates, samples, and tools that can be shared across plugins. The practical "
        "benefit of the unidirectional dependency structure (Skills \u2192 Agents \u2192 "
        "Gateway) lies in the independent replaceability of each layer. Adding or modifying "
        "agents does not affect the Gateway, and replacing the runtime does not require "
        "changes to Skills or Agents."
    )

    add_image_if_exists(
        doc,
        "fig_architecture.png",
        "Declarative Multi-Agent Plugin Architecture \u2014 Clean Architecture-based layer structure",
    )

    doc.add_paragraph()


def section_04_mechanism(doc):
    """Section 4: Operational Mechanism."""
    add_heading(doc, "Operational Mechanism \u2014 Prompt Assembly and Tier-Based Execution", level=1)

    # -- Prompt Assembly --
    add_heading(doc, "Prompt Assembly", level=2)

    add_body(
        doc,
        "When a delegation skill spawns an agent, the prompt is assembled in three layers. "
        "This ordering is designed to maximize the runtime's prefix cache hit rate."
    )

    add_body_bold(doc, "Layer 1 \u2014 Shared Static (runtime-mapping).")
    add_body(
        doc,
        "Rules applied universally to all agents. It includes mapping information from "
        "the Gateway's runtime-mapping.yaml that transforms abstract declarations into "
        "concrete models and tools. Since this is identical across all agent invocations "
        "within a session, cache efficiency is highest."
    )

    add_body_bold(doc, "Layer 2 \u2014 Per-Agent Static (AGENT.md + agentcard.yaml + tools.yaml).")
    add_body(
        doc,
        "Three files containing each agent's role, workflow, and tool definitions are "
        "synthesized into a single prompt. AGENT.md provides goals and methods (WHY+HOW), "
        "agentcard.yaml provides identity and constraints (WHO+WHAT+WHEN), and tools.yaml "
        "provides abstract tool interfaces. The cache is reused for repeated invocations "
        "of the same agent."
    )

    add_body_bold(doc, "Layer 3 \u2014 Dynamic (task instructions).")
    add_body(
        doc,
        "The specific task content delivered by the skill. This is the only layer that "
        "changes with every invocation, consisting of five items: task objective (TASK), "
        "expected deliverable (EXPECTED OUTCOME), mandatory requirements (MUST DO), "
        "prohibited actions (MUST NOT DO), and context (CONTEXT)."
    )

    add_body(
        doc,
        "The Gateway's runtime-mapping.yaml performs three transformations during this "
        "assembly process. Model concretization maps an agent's abstract tier (e.g., "
        "`tier: HIGH`) to an actual model (e.g., `claude-opus-4-6`). Tool concretization "
        "maps abstract tools from tools.yaml to actual tools. Forbidden action concretization "
        "maps forbidden action categories from agentcard.yaml to actual tools. The final "
        "set of tools provided to the agent is (concretized tools) - (forbidden tools)."
    )

    add_image_if_exists(
        doc,
        "fig_prompt_assembly.png",
        "Prompt Assembly Process \u2014 Three-layer synthesis of static and dynamic layers",
    )

    # -- 4-Tier Model Mapping --
    add_heading(doc, "4-Tier Model Mapping", level=2)

    add_body(
        doc,
        "Agent capability requirements are declared across four abstract tiers: HEAVY, HIGH, "
        "MEDIUM, and LOW. Each tier is assigned an optimal model based on cost-capability "
        "trade-offs, and agents of the same role are separated into tier variants to ensure "
        "cost efficiency."
    )

    make_table(
        doc,
        headers=["Tier", "Characteristics", "Suitable Tasks"],
        rows=[
            ["HEAVY", "Highest capability + large budget", "Extended reasoning, large-scale multi-file tasks"],
            ["HIGH", "Highest capability", "Complex decision-making, deep analysis"],
            ["MEDIUM", "Balanced", "Feature implementation, general analysis"],
            ["LOW", "Fast and low-cost", "Single lookups, simple modifications"],
        ],
        col_widths=[3, 5, 8],
    )

    add_body(
        doc,
        "A noteworthy aspect is the escalation mechanism. When a LOW-tier agent recognizes "
        "its own limitations, it stops work and reports an escalation to a higher tier. This "
        "self-awareness-based step-by-step delegation is a key mechanism that prevents "
        "resource waste while ensuring quality. Tier variant agents inherit the base agent's "
        "configuration and only describe the parts that need overriding, minimizing "
        "duplication."
    )

    add_image_if_exists(
        doc,
        "fig_tier_mapping.png",
        "4-Tier Model Mapping \u2014 Transformation from abstract tier declaration to concrete LLM model",
    )

    doc.add_paragraph()


def section_05_orchestration(doc):
    """Section 5: Execution Flow."""
    add_heading(doc, "Execution Flow \u2014 Skill-Based Orchestration", level=1)

    add_body(doc, "Skill activation occurs through two pathways.")

    add_body_bold(doc, "Direct Activation")
    add_body(
        doc,
        "When the runtime detects an explicit slash command (`/plugin:skill`) or a natural "
        "language request that shows high similarity to a skill's metadata (frontmatter "
        "description), the corresponding skill is loaded immediately. This is a path where "
        "the user's intent is directly linked to a specific skill without intermediate layers."
    )

    add_body_bold(doc, "Core-Mediated Activation")
    add_body(
        doc,
        "When a user's request is ambiguous or spans multiple skills, the runtime first "
        "activates the Core skill. The Core skill is limited to determining request intent "
        "and routing to the appropriate skill. Actual task execution is handled by the "
        "routed target skill (Orchestrator, Planning, etc.)."
    )

    add_body(
        doc,
        "Execution paths diverge based on skill type. Delegation skills (Core, Planning, "
        "Orchestrator) invoke agents to delegate tasks, while direct skills (Setup, Utility) "
        "use Gateway tools directly. This dual-path design is the result of simultaneously "
        "applying Clean Architecture's separation of concerns and the YAGNI principle. "
        "Routing procedural tasks like installation or status checks through agents would "
        "generate unnecessary LLM calls, making it rational for direct skills to access "
        "the Gateway directly."
    )

    add_body(
        doc,
        "At session start, the runtime scans the `skills/` directory to automatically "
        "discover all skills. Skills become immediately available simply by placing them "
        "in the directory, with no separate registration procedure required."
    )

    add_image_if_exists(
        doc,
        "fig_skill_activation.png",
        "Skill Activation Pathways \u2014 Dual paths of direct activation and Core-mediated activation",
    )

    doc.add_paragraph()


def section_06_case_study(doc):
    """Section 6: Implementation Case Study."""
    add_heading(doc, "Implementation Case Study \u2014 DMAP and the abra Plugin", level=1)

    add_body(
        doc,
        "The practical implementation of the declarative architecture described above is "
        "DMAP (Declarative Multi-Agent Plugin). DMAP is an open-source plugin architecture "
        "standard that defines multi-agent systems using only Markdown and YAML without "
        "any code. The DMAP Builder is a tool that automatically generates plugins "
        "according to this standard."
    )

    add_body(
        doc,
        "The abra plugin serves as an application case for DMAP. abra is a plugin that "
        "automatically generates Dify workflow DSL from a single natural language command "
        "and performs prototyping. Five agents (scenario-analyst, dsl-architect, "
        "plan-writer, prototype-runner, agent-developer) each handle their specialized "
        "roles, while eight skills orchestrate the workflow from scenario analysis to "
        "code development. The entire structure is composed solely of Markdown and YAML "
        "\u2014 no programming code is used in any aspect, including agent role definitions, "
        "tool connections, or execution ordering."
    )

    add_body(
        doc,
        "This implementation demonstrates that the declarative approach is not merely a "
        "theoretical possibility but a viable, working system. The entire project is "
        "publicly available on GitHub (https://github.com/unicorn-plugins/dmap)."
    )

    add_image_if_exists(
        doc,
        "fig_abra_example.png",
        "abra Plugin Structure \u2014 Declarative composition of 5 agents and 8 skills",
    )

    doc.add_paragraph()


def section_07_outlook(doc):
    """Section 7: Outlook."""
    add_heading(doc, "Outlook \u2014 The Future of the Declarative Agent Ecosystem", level=1)

    add_body(
        doc,
        "The impact of the declarative approach on the agent development ecosystem is "
        "noteworthy across three dimensions."
    )

    add_body(
        doc,
        "First, it changes the collaboration model between developers and non-developers. "
        "When agent system design shifts from code to declarations, a division of labor "
        "becomes possible where domain experts directly define agent roles and constraints "
        "while developers handle infrastructure mapping. This expands the range of "
        "participants in agent system design."
    )

    add_body(
        doc,
        "Second, runtime portability opens the possibility of agent reuse. If a single "
        "plugin can operate across various runtimes such as Claude Code, Codex CLI, and "
        "Gemini CLI simply by swapping runtime-mapping.yaml, the agent ecosystem can "
        "evolve in a direction free from platform lock-in."
    )

    add_body(
        doc,
        "Third, there is the direction of agent system standardization. By formalizing "
        "agent roles, capabilities, constraints, and handoffs on top of universal formats "
        "like Markdown and YAML, a foundation is established for systematic verification, "
        "sharing, and composition of agent packages."
    )

    add_body(
        doc,
        "The paradigm shift from code to declaration ultimately separates the concerns "
        "of \"what to build\" from \"how to execute it.\" This suggests that the principle "
        "software engineering has repeatedly proven \u2014 that raising the level of "
        "abstraction makes system complexity manageable \u2014 holds true in the domain "
        "of LLM agents as well."
    )

    add_note(
        doc,
        "DMAP standard documents and samples are available on GitHub: "
        "https://github.com/unicorn-plugins/dmap",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Document margins (2.5cm all sides)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font settings (Arial, 11pt, 1.5 line spacing)
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Title (22pt, Bold, Center)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Declarative Multi-Agent Plugin: Designing Agent Systems Without Code")
    run.font.size = Pt(22)
    run.font.name = FONT_NAME
    run.bold = True

    # Subtitle (13pt, color #555555, Center)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Redefining LLM Agent Orchestration with Clean Architecture Principles")
    run.font.size = Pt(13)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(
        int(COLOR_SUBTITLE[0:2], 16),
        int(COLOR_SUBTITLE[2:4], 16),
        int(COLOR_SUBTITLE[4:6], 16)
    )

    doc.add_paragraph()  # blank line

    # 7 sections in sequence
    section_01_intro(doc)
    section_02_declarative(doc)
    section_03_architecture(doc)
    section_04_mechanism(doc)
    section_05_orchestration(doc)
    section_06_case_study(doc)
    section_07_outlook(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Medium article generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
