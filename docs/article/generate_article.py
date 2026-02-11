#!/usr/bin/env python3
"""
DMAP 빌더 소개 블로그 아티클 생성 스크립트
python-docx 기반으로 Word(.docx) 형식의 블로그 아티클을 생성함.

사용법: python docs/article/generate_article.py
출력: docs/article/DMAP-builder-article.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "DMAP-builder-article.docx")

# ---------------------------------------------------------------------------
# Font settings
# ---------------------------------------------------------------------------
FONT_NAME = "맑은 고딕"
FONT_FALLBACK = "Arial"


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """테이블 테두리 추가."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}></w:tblPr>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None):
    """포맷된 테이블 생성 (파란 헤더 배경 #D5E8F0)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D5E8F0")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = FONT_NAME

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_image_if_exists(doc, filename, caption, width_inches=5.5):
    """이미지 삽입 (파일 없으면 경고 출력 후 graceful skip)."""
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        print(f"  [WARNING] 이미지 누락: {filename}")
        p = doc.add_paragraph()
        run = p.add_run(f"[그림 누락: {filename}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run.italic = True


def add_heading(doc, text, level=1):
    """제목 추가 (맑은 고딕, H1=18pt, H2=14pt)."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return h


def add_body(doc, text):
    """본문 추가 (맑은 고딕, 11pt, 행간 1.5줄)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


def add_body_bold(doc, text):
    """볼드 본문 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    return p


def add_quote(doc, text, label="TIP"):
    """TIP 박스 (1x1 테이블, 배경 #E8F0FE, 왼쪽 굵은 파란 테두리 #3366CC)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
    run.font.name = FONT_NAME
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    set_cell_shading(cell, "E8F0FE")
    # 왼쪽 굵은 파란 테두리
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:left w:val="single" w:sz="24" w:color="3366CC"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)


def add_bullet(doc, text):
    """불릿 리스트 추가."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


# ---------------------------------------------------------------------------
# Section functions
# ---------------------------------------------------------------------------

def section_01_problem(doc):
    """섹션 1: 멀티에이전트 개발의 불편함."""
    add_heading(doc, "멀티에이전트 개발, 지금 뭐가 불편한가요?", level=1)

    add_body(doc,
        "오케스트라를 상상해보세요. 지휘자가 악보 없이 수십 명의 연주자에게 "
        "\"알아서 잘 맞춰봐\"라고 한다면 어떨까요? 아마 끔찍한 소음이 나겠죠. "
        "AI 에이전트도 마찬가지입니다."
    )
    add_body(doc,
        "에이전트 1개일 때는 괜찮습니다. 하지만 여러 에이전트가 협업해야 하는 순간, "
        "복잡도가 폭발적으로 늘어나요. 기존 방식은 Python SDK로 직접 코딩하거나, "
        "LangChain, CrewAI, AutoGen 같은 프레임워크를 배워야 합니다. "
        "프레임워크마다 문법이 다르고, 특정 런타임에 종속되어 버리죠."
    )
    add_body(doc,
        "진입장벽은 높고, 유지보수는 어렵고, 다른 환경으로 옮기려면 처음부터 다시 만들어야 합니다. "
        "여기서 근본적인 질문이 떠오릅니다. "
        "\"에이전트를 정의하는 데 꼭 코드가 필요할까?\""
    )


def section_02_overview(doc):
    """섹션 2: DMAP 빌더 소개."""
    add_heading(doc, "DMAP 빌더란?", level=1)

    add_body(doc,
        "레고를 떠올려보세요. 설명서만 보고 블록을 조립하면 멋진 작품이 완성되잖아요? "
        "DMAP 빌더도 같은 원리입니다. 코드 한 줄 없이, Markdown과 YAML 파일만 작성하면 "
        "멀티에이전트 플러그인이 뚝딱 만들어집니다."
    )
    add_body(doc,
        "DMAP는 Declarative Multi-Agent Plugin의 약자입니다. "
        "\"선언형\"이라는 말이 어렵게 느껴질 수 있는데, 쉽게 말하면 "
        "\"어떻게(HOW)\" 대신 \"무엇을(WHAT)\"만 적는 방식이에요. "
        "택시 기사에게 경로를 일일이 지시하는 대신, \"강남역이요\"라고 목적지만 말하는 것처럼요."
    )

    add_body_bold(doc, "DMAP의 핵심 가치")
    add_bullet(doc, "선언형 명세 — 코드 대신 문서로 에이전트를 정의합니다")
    add_bullet(doc, "런타임 중립 — Claude Code, Codex CLI 등 어디서든 동작해요")
    add_bullet(doc, "관심사 분리 — 역할별로 깔끔하게 나뉘어 유지보수가 쉽습니다")
    add_bullet(doc, "비개발자 접근성 — 마크다운만 쓸 줄 알면 누구나 만들 수 있어요")
    add_bullet(doc, "도메인 범용 — 어떤 업무 영역이든 적용 가능합니다")

    doc.add_paragraph()  # 간격

    make_table(doc,
        headers=["구분", "기존 프레임워크", "DMAP"],
        rows=[
            ["명세 방식", "Python/TypeScript 코드", "Markdown + YAML"],
            ["런타임", "프레임워크 전용 엔진", "Claude Code, Codex CLI 등"],
            ["진입장벽", "코딩 필수", "마크다운 작성 가능하면 OK"],
            ["에이전트 교체", "코드 수정 필요", "YAML 한 줄 변경"],
        ],
        col_widths=[4, 6, 6],
    )


def section_03_core_concepts(doc):
    """섹션 3: 핵심 개념 설명."""
    add_heading(doc, "핵심 개념: 회사에 비유하면", level=1)

    add_body(doc,
        "DMAP의 구조를 회사 조직에 비유하면 이해하기 쉬워요. "
        "회사에는 부서장, 전문가, 통역사가 있잖아요? DMAP도 비슷한 구조입니다."
    )

    add_body_bold(doc, "Skills = 부서장")
    add_body(doc,
        "스킬은 일을 배분하는 부서장 같은 역할이에요. "
        "사용자의 요청을 받아서 어떤 에이전트에게 어떤 일을 시킬지 결정합니다. "
        "SKILL.md 파일 하나로 워크플로우 전체를 선언하죠."
    )

    add_body_bold(doc, "Agents = 전문가")
    add_body(doc,
        "에이전트는 실제 일을 하는 전문가입니다. "
        "각 에이전트는 AGENT.md(역할 정의), agentcard.yaml(메타데이터), "
        "tools.yaml(사용 가능한 도구) 3개 파일로 구성돼요."
    )

    add_body_bold(doc, "Gateway = 통역사")
    add_body(doc,
        "게이트웨이는 추상적인 선언을 구체적인 실행 환경으로 번역하는 통역사입니다. "
        "에이전트가 \"파일 검색 도구\"라고 선언하면, 게이트웨이가 실제 런타임에서 "
        "어떤 도구를 쓸지 매핑해줘요."
    )

    add_body_bold(doc, "리소스 마켓플레이스 = 사내 공유 드라이브")
    add_body(doc,
        "가이드, 템플릿, 샘플, 도구를 모아둔 공유 저장소입니다. "
        "여러 플러그인이 함께 사용할 수 있는 공용 자원이에요."
    )

    add_body(doc,
        "이 컴포넌트들은 Clean Architecture 원칙을 따릅니다. "
        "Skills → Agents → Gateway 순서로 단방향 의존하기 때문에, "
        "각 부분을 독립적으로 수정하거나 교체할 수 있어요."
    )

    add_quote(doc,
        "선언형이란? \"어떻게(HOW)\" 대신 \"무엇을(WHAT)\"만 적으면 됩니다! "
        "SQL처럼 원하는 결과만 말하면, 실행은 엔진이 알아서 해줘요."
    )

    doc.add_paragraph()  # 간격
    add_image_if_exists(doc, "fig_architecture.png",
        "DMAP 아키텍처 — Clean Architecture 기반 컴포넌트 구성")


def section_04_how_it_works(doc):
    """섹션 4: 동작 원리."""
    add_heading(doc, "동작 원리: 어떻게 돌아갈까요?", level=1)

    # --- 프롬프트 조립 ---
    add_heading(doc, "프롬프트 조립 — 여행 가이드북 만들기", level=2)
    add_body(doc,
        "해외여행을 갈 때 가이드북을 만든다고 생각해보세요. "
        "기본 정보(공통 규칙)에 목적지별 정보(에이전트 역할)를 합치고, "
        "마지막에 오늘의 일정(작업 지시)을 추가하면 완성되죠?"
    )
    add_body(doc,
        "DMAP도 같은 방식으로 프롬프트를 조립합니다. "
        "먼저 Gateway의 runtime-mapping으로 공통 설정을 로드하고, "
        "에이전트 패키지 3개 파일(AGENT.md, agentcard.yaml, tools.yaml)을 합친 뒤, "
        "마지막으로 스킬이 전달하는 동적 작업 지시를 붙여서 최종 프롬프트를 완성합니다."
    )
    add_image_if_exists(doc, "fig_prompt_assembly.png",
        "프롬프트 조립 과정 — 3단계 레이어 구성")

    # --- 4-Tier 매핑 ---
    add_heading(doc, "4-Tier 매핑 — 택시 vs 버스 vs KTX", level=2)
    add_body(doc,
        "모든 일에 택시를 탈 필요는 없잖아요? "
        "가까운 곳은 버스로, 먼 곳은 KTX로 가면 비용도 절약되고 효율적이죠. "
        "DMAP의 4-Tier 매핑도 같은 원리입니다."
    )
    add_body(doc,
        "작업의 복잡도에 따라 HEAVY, HIGH, MEDIUM, LOW 4단계로 나누고, "
        "각 단계에 적합한 AI 모델을 배정해요. "
        "간단한 파일 검색은 가벼운 모델(Haiku)이, "
        "복잡한 아키텍처 설계는 강력한 모델(Opus)이 담당합니다. "
        "작업 중 난이도가 올라가면 자동으로 상위 모델로 에스컬레이션되기도 해요."
    )
    add_image_if_exists(doc, "fig_tier_mapping.png",
        "4-Tier 모델 매핑 — 복잡도별 최적 모델 배정")

    # --- 스킬 활성화 ---
    add_heading(doc, "스킬 활성화 — 114 안내교환대", level=2)
    add_body(doc,
        "예전 114 안내교환대를 기억하시나요? "
        "\"피자집 전화번호 알려주세요\"라고 하면 알아서 연결해주었죠. "
        "DMAP의 스킬 활성화도 비슷합니다."
    )
    add_body(doc,
        "두 가지 경로가 있어요. "
        "첫째, 슬래시 명령어로 직접 호출하는 방법 — 정확히 어떤 스킬이 필요한지 알 때 사용합니다. "
        "둘째, Core 스킬을 경유하는 방법 — 모호한 요청이 들어오면 Core가 "
        "의도를 분석해서 적절한 스킬로 라우팅해줘요. "
        "런타임은 skills/ 디렉토리를 자동 스캔해서 사용 가능한 스킬을 발견합니다."
    )
    add_image_if_exists(doc, "fig_skill_activation.png",
        "스킬 활성화 경로 — 직접 호출과 Core 경유")


def section_05_dev_method(doc):
    """섹션 5: 플러그인 개발 방법."""
    add_heading(doc, "플러그인 개발 방법: 요리 레시피 따라하기", level=1)

    add_body(doc,
        "요리를 할 때 레시피를 따라하면 누구나 맛있는 음식을 만들 수 있잖아요? "
        "DMAP 플러그인 개발도 마찬가지입니다. "
        "4단계 레시피만 따라가면 플러그인이 완성돼요."
    )

    doc.add_paragraph()  # 간격

    make_table(doc,
        headers=["Phase", "단계", "비유"],
        rows=[
            ["1", "요구사항 수집", "어떤 요리를 만들지 정하기"],
            ["2", "설계 및 계획", "레시피 작성하기"],
            ["3", "플러그인 개발", "실제 요리하기"],
            ["4", "검증 및 완료", "시식하고 플레이팅"],
        ],
        col_widths=[3, 5, 8],
    )

    doc.add_paragraph()  # 간격

    add_body(doc,
        "각 Phase가 끝날 때마다 사용자의 확인을 받습니다. "
        "\"이 요리 맞죠?\"라고 중간중간 물어보는 셈이에요. "
        "덕분에 엉뚱한 결과물이 나오는 걸 미리 방지할 수 있습니다."
    )


def section_06_example(doc):
    """섹션 6: 실전 예제 (abra 플러그인)."""
    add_heading(doc, "실전 예제: abra 플러그인", level=1)

    add_body(doc,
        "백문이 불여일견! 실제로 DMAP으로 만든 플러그인을 살펴볼게요. "
        "\"abra\"는 Dify 워크플로우 DSL 자동화 플러그인입니다. "
        "Dify라는 AI 워크플로우 플랫폼에서 사용하는 DSL(Domain Specific Language)을 "
        "자동으로 생성하고, 프로토타이핑까지 해주는 똑똑한 플러그인이에요."
    )

    add_body_bold(doc, "에이전트 5명의 전문가 팀")
    add_bullet(doc, "scenario-analyst — 요구사항을 분석하고 시나리오를 만드는 기획자")
    add_bullet(doc, "dsl-architect — Dify DSL 코드를 설계하는 아키텍트")
    add_bullet(doc, "plan-writer — 개발계획서를 작성하는 PM")
    add_bullet(doc, "prototype-runner — Dify에서 프로토타이핑을 실행하는 테스터")
    add_bullet(doc, "agent-developer — 최종 Agent 코드를 개발하는 개발자")

    add_body_bold(doc, "스킬 8개 — 부서장들의 업무 지시서")
    add_body(doc,
        "setup, core, dify-setup, scenario, dsl-generate, prototype, dev-plan, develop "
        "— 각 스킬이 워크플로우의 한 단계를 담당합니다."
    )

    add_body(doc,
        "사용자 입장에서는 시나리오 생성 → DSL 생성 → 프로토타이핑 → "
        "개발계획 → Agent 개발 순서로 자연스럽게 진행돼요. "
        "복잡한 내부 구조를 몰라도 됩니다!"
    )

    doc.add_paragraph()  # 간격
    add_image_if_exists(doc, "fig_abra_example.png",
        "abra 플러그인 구조 — 5개 에이전트와 8개 스킬의 협업")


def section_07_roadmap(doc):
    """섹션 7: 로드맵."""
    add_heading(doc, "앞으로의 로드맵", level=1)

    add_body(doc,
        "DMAP는 계속 진화하고 있습니다. 앞으로 어떤 모습이 될지 살짝 엿볼게요."
    )

    add_body_bold(doc, "현재")
    add_bullet(doc, "Claude Code 런타임 기반으로 안정적으로 동작 중")

    add_body_bold(doc, "단기 계획")
    add_bullet(doc, "다중 런타임 지원 — Codex CLI, Gemini CLI 등 다양한 환경에서 실행")

    add_body_bold(doc, "중기 계획")
    add_bullet(doc, "플러그인 마켓플레이스 — 만든 플러그인을 공유하고 다운로드하는 생태계 구축")
    add_bullet(doc, "커뮤니티 생태계 — 개발자들이 함께 성장하는 오픈소스 커뮤니티")

    add_body_bold(doc, "장기 비전")
    add_bullet(doc, "비개발자용 시각적 플러그인 빌더 — 드래그 앤 드롭으로 플러그인을 만드는 노코드 UI")


def section_08_quickstart(doc):
    """섹션 8: 빠른 시작 가이드."""
    add_heading(doc, "지금 바로 시작하기", level=1)

    add_body(doc,
        "어렵게 느껴지셨나요? 걱정 마세요. 4단계면 시작할 수 있습니다!"
    )

    add_bullet(doc, "Step 1: DMAP 표준 문서를 읽고 전체 구조를 파악하세요")
    add_bullet(doc, "Step 2: samples/abra 예제를 참고해서 나만의 플러그인 폴더를 만드세요")
    add_bullet(doc, "Step 3: AGENT.md, agentcard.yaml, tools.yaml을 작성하세요")
    add_bullet(doc, "Step 4: 스킬과 게이트웨이를 연결하고 테스트하세요")

    doc.add_paragraph()  # 간격

    add_body(doc,
        "코드 한 줄 없이 멀티에이전트 플러그인을 만들 수 있다는 걸 "
        "직접 경험해보세요. 지금 바로 시작해보세요!"
    )

    add_quote(doc,
        "DMAP 표준 문서와 샘플은 GitHub에서 확인하세요: "
        "https://github.com/unicorn-plugins/dmap"
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # 문서 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 타이틀
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DMAP 빌더: 코드 없이 멀티에이전트 플러그인 만들기")
    run.font.size = Pt(22)
    run.font.name = FONT_NAME
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("선언형 멀티에이전트 플러그인의 세계로 초대합니다")
    run.font.size = Pt(13)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 빈 줄

    # 각 섹션 순차 호출
    section_01_problem(doc)
    section_02_overview(doc)
    section_03_core_concepts(doc)
    section_04_how_it_works(doc)
    section_05_dev_method(doc)
    section_06_example(doc)
    section_07_roadmap(doc)
    section_08_quickstart(doc)

    doc.save(OUTPUT_PATH)
    print(f"Article generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
