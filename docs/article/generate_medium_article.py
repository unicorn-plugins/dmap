#!/usr/bin/env python3
"""
DMAP Medium 아티클 생성 스크립트
python-docx 기반으로 Word(.docx) 형식의 Medium 아티클을 생성함.

사용법: python docs/article/generate_medium_article.py
출력: docs/article/DMAP-medium-article.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "DMAP-medium-article.docx")

# ---------------------------------------------------------------------------
# Style Constants
# ---------------------------------------------------------------------------
FONT_NAME = "맑은 고딕"
FONT_FALLBACK = "Arial"

# Colors
COLOR_HEADING_H1 = "1A1A1A"
COLOR_HEADING_H2 = "2D2D2D"
COLOR_SUBTITLE = "555555"
COLOR_TABLE_HEADER_BG = "D5E8F0"
COLOR_TABLE_BORDER = "CCCCCC"
COLOR_NOTE_BG = "F0F4F8"
COLOR_NOTE_BORDER = "4A6FA5"
COLOR_ACCENT = "2C5282"


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """테이블 테두리 추가 (색상: #CCCCCC)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}></w:tblPr>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None):
    """포맷된 테이블 생성 (파란 헤더 배경 #D5E8F0)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = FONT_NAME

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_image_if_exists(doc, filename, caption, width_inches=5.5):
    """이미지 삽입 (파일 없으면 경고 출력 후 graceful skip)."""
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        print(f"  [WARNING] 이미지 누락: {filename}")
        p = doc.add_paragraph()
        run = p.add_run(f"[그림 누락: {filename}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run.italic = True


def add_heading(doc, text, level=1):
    """제목 추가 (H1=16pt 색상 #1A1A1A, H2=13pt 색상 #2D2D2D)."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(
                int(COLOR_HEADING_H1[0:2], 16),
                int(COLOR_HEADING_H1[2:4], 16),
                int(COLOR_HEADING_H1[4:6], 16)
            )
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(
                int(COLOR_HEADING_H2[0:2], 16),
                int(COLOR_HEADING_H2[2:4], 16),
                int(COLOR_HEADING_H2[4:6], 16)
            )
    return h


def add_body(doc, text):
    """본문 추가 (맑은 고딕, 11pt, 행간 1.5줄)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


def add_body_bold(doc, text):
    """볼드 본문 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    return p


def add_note(doc, text, label="NOTE"):
    """NOTE 박스 (1x1 테이블, 배경 #F0F4F8, 왼쪽 굵은 테두리 #4A6FA5 두께 24pt)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(
        int(COLOR_NOTE_BORDER[0:2], 16),
        int(COLOR_NOTE_BORDER[2:4], 16),
        int(COLOR_NOTE_BORDER[4:6], 16)
    )
    run.font.name = FONT_NAME
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    set_cell_shading(cell, COLOR_NOTE_BG)
    # 왼쪽 굵은 테두리
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:left w:val="single" w:sz="24" w:color="{COLOR_NOTE_BORDER}"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)


def add_bullet(doc, text):
    """불릿 리스트 추가."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.5
    return p


# ---------------------------------------------------------------------------
# Section functions
# ---------------------------------------------------------------------------

def section_01_intro(doc):
    """섹션 1: 도입."""
    add_heading(doc, "멀티에이전트 시스템의 구조적 과제", level=1)

    add_body(
        doc,
        "LLM 에이전트가 단일 개체에서 다수의 협업 시스템으로 확장될 때, 복잡도는 선형이 아닌 "
        "조합적으로 증가한다. 에이전트 간 역할 경계, 도구 접근 권한, 실행 순서, 에스컬레이션 조건 등 "
        "관리해야 할 관계의 수가 에이전트 수의 제곱에 비례하여 늘어나는 구조적 문제가 발생한다."
    )

    add_body(
        doc,
        "기존 접근법인 LangChain, CrewAI, AutoGen 등의 프레임워크는 이 문제를 "
        "Python/TypeScript SDK 코드로 해결한다. 에이전트의 정의, 도구 연결, 오케스트레이션 로직이 "
        "모두 프로그래밍 언어 안에 혼재되며, 특정 SDK에 강결합되는 구조를 형성한다. 이로 인해 런타임 "
        "환경이 변경되면 전체 시스템을 재작성해야 하고, 도메인 전문가나 비개발자는 에이전트 시스템의 "
        "설계에 참여할 수 없다."
    )

    add_body(
        doc,
        "근본적 질문이 제기된다. 에이전트의 역할과 관계를 정의하는 데 반드시 코드가 필요한가. "
        "에이전트가 \"무엇을 할 수 있는가\"를 선언하는 것과 \"어떻게 동작하는가\"를 코딩하는 것은 "
        "본질적으로 다른 행위이다."
    )

    doc.add_paragraph()


def section_02_declarative(doc):
    """섹션 2: 선언형 접근."""
    add_heading(doc, "선언형 접근의 본질 \u2014 \"What\" 대 \"How\"", level=1)

    add_body(
        doc,
        "선언형 프로그래밍은 \"원하는 결과\"를 기술하고 \"달성 방법\"은 실행 환경에 위임하는 "
        "패러다임이다. SQL은 데이터의 조건을 선언하되 탐색 알고리즘을 지정하지 않고, HTML/CSS는 "
        "문서의 구조와 스타일을 선언하되 렌더링 엔진의 구현을 강제하지 않으며, Kubernetes YAML은 "
        "원하는 인프라 상태를 선언하되 컨테이너 스케줄링 방식을 지시하지 않는다. 이들의 공통점은 "
        "실행 환경(런타임)과 선언(명세)의 분리에 있다."
    )

    add_body(
        doc,
        "멀티에이전트 시스템에도 동일한 원리가 적용 가능하다. 에이전트의 역할, 역량, 제약을 "
        "Markdown과 YAML로 선언하고, 실행은 런타임(Claude Code, Codex CLI, Gemini CLI 등)에 "
        "위임하는 방식이다. 이 접근의 핵심 가치는 두 가지로 요약된다. 첫째, 런타임 중립성 \u2014 "
        "`tier: HIGH` 같은 추상 선언은 어떤 런타임에서든 해당 환경의 최적 모델로 해석 가능하다. "
        "둘째, 비개발자 접근성 \u2014 Markdown과 YAML만 알면 도메인 전문가도 에이전트 시스템을 "
        "구축할 수 있다."
    )

    make_table(
        doc,
        headers=["비교 항목", "기존 프레임워크", "선언형 접근"],
        rows=[
            ["에이전트 정의", "Python/TypeScript SDK 코드", "Markdown + YAML"],
            ["오케스트레이션", "그래프/함수 호출 체인 코드", "스킬 프롬프트(자연어)"],
            ["런타임 종속성", "특정 SDK에 강결합", "런타임 중립"],
            ["아키텍처 원칙", "없거나 프레임워크 의존", "Clean Architecture"],
            ["도구 연결", "코드에서 도구 객체 생성", "추상 선언(tools.yaml) \u2192 Gateway 매핑"],
            ["티어 관리", "없음", "4-Tier(HEAVY/HIGH/MEDIUM/LOW) + 런타임 자동 매핑"],
            ["이식성", "낮음 (전체 재작성)", "높음 (runtime-mapping.yaml만 교체)"],
        ],
        col_widths=[4, 6, 6],
    )

    doc.add_paragraph()


def section_03_architecture(doc):
    """섹션 3: 아키텍처."""
    add_heading(doc, "아키텍처 설계 \u2014 Clean Architecture의 에이전트 적용", level=1)

    add_body(
        doc,
        "Clean Architecture의 핵심 원칙인 단방향 의존과 관심사 분리는 소프트웨어 시스템의 "
        "장기적 유지보수성을 보장하는 검증된 설계 철학이다. 선언형 멀티에이전트 아키텍처는 이 원칙을 "
        "에이전트 시스템에 직접 적용하여 3개 핵심 계층으로 구성된다."
    )

    add_body_bold(doc, "Skills (Controller 계층)")
    add_body(
        doc,
        "사용자 요청의 진입점이자 오케스트레이션 계층이다. 요청의 의도를 분류하고 적절한 에이전트에 "
        "작업을 위임하는 라우팅 역할을 수행한다. 위임형 스킬(Core, Planning, Orchestrator)은 "
        "에이전트를 호출하고, 직결형 스킬(Setup, Utility)은 Gateway를 직접 사용한다. 실행 경로에 "
        "따라 불필요한 계층을 생략하는 YAGNI 원칙이 적용된다."
    )

    add_body_bold(doc, "Agents (Service 계층)")
    add_body(
        doc,
        "역할 기반의 자율 실행 단위이다. 각 에이전트는 AGENT.md(목표와 워크플로우 프롬프트)와 "
        "agentcard.yaml(정체성, 역량, 제약 선언)의 쌍으로 구성된다. 하나의 에이전트는 하나의 "
        "전문 역할에만 집중하며, 역할 경계를 넘는 작업은 핸드오프 규칙에 따라 다른 에이전트에 위임한다."
    )

    add_body_bold(doc, "Gateway (Infrastructure 계층)")
    add_body(
        doc,
        "추상 선언과 구체 런타임 간의 매핑 테이블을 제공하는 인프라 계층이다. runtime-mapping.yaml이 "
        "에이전트의 추상 티어를 실제 LLM 모델로, 추상 도구를 실제 도구 구현체로 변환한다. 런타임 "
        "환경이 변경되면 이 매핑 파일만 교체하면 된다."
    )

    add_body(
        doc,
        "이 3개 계층 외에 리소스 마켓플레이스가 플러그인 간 공유 가능한 가이드, 템플릿, 샘플, "
        "도구의 공용 풀로 기능한다. 단방향 의존 구조(Skills \u2192 Agents \u2192 Gateway)의 "
        "실질적 이점은 각 계층의 독립적 교체 가능성에 있다. 에이전트를 추가하거나 변경해도 "
        "Gateway에 영향을 주지 않고, 런타임을 교체해도 Skills와 Agents는 수정 없이 동작한다."
    )

    add_image_if_exists(
        doc,
        "fig_architecture.png",
        "선언형 멀티에이전트 플러그인 아키텍처 \u2014 Clean Architecture 기반 계층 구조",
    )

    doc.add_paragraph()


def section_04_mechanism(doc):
    """섹션 4: 동작 메커니즘."""
    add_heading(doc, "동작 메커니즘 \u2014 프롬프트 조립과 티어 기반 실행", level=1)

    # -- 프롬프트 조립 --
    add_heading(doc, "프롬프트 조립 (Prompt Assembly)", level=2)

    add_body(
        doc,
        "위임형 스킬이 에이전트를 스폰할 때, 3단계 레이어로 프롬프트가 조립된다. "
        "이 순서는 런타임의 prefix 캐시 적중률을 극대화하기 위해 설계된 것이다."
    )

    add_body_bold(doc, "1단계 \u2014 공통 정적 (runtime-mapping).")
    add_body(
        doc,
        "모든 에이전트에 공통으로 적용되는 규칙이다. Gateway의 runtime-mapping.yaml에서 추상 선언을 "
        "구체 모델과 도구로 변환하는 매핑 정보가 포함된다. 세션 내 모든 에이전트 호출에서 동일하므로 "
        "캐시 효율이 가장 높다."
    )

    add_body_bold(doc, "2단계 \u2014 에이전트별 정적 (AGENT.md + agentcard.yaml + tools.yaml).")
    add_body(
        doc,
        "각 에이전트의 역할, 워크플로우, 도구 정의를 포함하는 3개 파일이 하나의 프롬프트로 합성된다. "
        "AGENT.md가 목표와 수행 방법(WHY+HOW)을, agentcard.yaml이 정체성과 제약(WHO+WHAT+WHEN)을, "
        "tools.yaml이 추상 도구 인터페이스를 각각 담당한다. 동일 에이전트의 반복 호출 시 캐시가 재활용된다."
    )

    add_body_bold(doc, "3단계 \u2014 동적 (작업 지시).")
    add_body(
        doc,
        "스킬이 전달하는 구체적 작업 내용이다. 매 호출마다 변경되는 유일한 레이어로, 작업 목표(TASK), "
        "기대 산출물(EXPECTED OUTCOME), 필수 요건(MUST DO), 금지 사항(MUST NOT DO), "
        "컨텍스트(CONTEXT)의 5항목으로 구성된다."
    )

    add_body(
        doc,
        "Gateway의 runtime-mapping.yaml은 이 조립 과정에서 세 가지 변환을 수행한다. 에이전트의 "
        "추상 티어(예: `tier: HIGH`)를 실제 모델(예: `claude-opus-4-6`)로 매핑하는 모델 구체화, "
        "tools.yaml의 추상 도구를 실제 도구로 매핑하는 툴 구체화, agentcard.yaml의 금지 액션 "
        "카테고리를 실제 도구로 매핑하는 금지액션 구체화가 그것이다. 최종적으로 에이전트에 제공되는 "
        "도구는 (구체화된 도구) - (금지 도구)의 결과이다."
    )

    add_image_if_exists(
        doc,
        "fig_prompt_assembly.png",
        "프롬프트 조립 과정 \u2014 정적 레이어와 동적 레이어의 3단계 합성",
    )

    # -- 4-Tier 모델 매핑 --
    add_heading(doc, "4-Tier 모델 매핑", level=2)

    add_body(
        doc,
        "에이전트의 역량 요구 수준은 HEAVY, HIGH, MEDIUM, LOW의 4단계 추상 티어로 선언된다. "
        "각 티어는 비용-역량 트레이드오프에 따라 최적 모델이 배정되며, 동일 역할의 에이전트를 "
        "티어별 변형으로 분리하여 비용 효율성을 확보한다."
    )

    make_table(
        doc,
        headers=["티어", "특성", "적합 작업"],
        rows=[
            ["HEAVY", "최고 역량 + 대규모 예산", "장시간 추론, 대규모 멀티파일 작업"],
            ["HIGH", "최고 역량", "복잡한 의사결정, 심층 분석"],
            ["MEDIUM", "균형", "기능 구현, 일반 분석"],
            ["LOW", "빠르고 저비용", "단건 조회, 간단한 수정"],
        ],
        col_widths=[3, 5, 8],
    )

    add_body(
        doc,
        "주목할 점은 에스컬레이션 메커니즘이다. LOW 티어 에이전트가 자기 한계를 인식하면 작업을 "
        "중단하고 상위 티어로의 에스컬레이션을 보고한다. 이 자기 인식 기반의 단계적 위임은 자원 "
        "낭비를 방지하면서 품질을 보장하는 핵심 메커니즘이다. 티어 변형 에이전트는 기본 에이전트의 "
        "설정을 상속(inherits)받고 오버라이드가 필요한 부분만 기술하므로 중복이 최소화된다."
    )

    add_image_if_exists(
        doc,
        "fig_tier_mapping.png",
        "4-Tier 모델 매핑 \u2014 추상 티어 선언에서 구체 LLM 모델로의 변환",
    )

    doc.add_paragraph()


def section_05_orchestration(doc):
    """섹션 5: 실행 흐름."""
    add_heading(doc, "실행 흐름 \u2014 스킬 기반 오케스트레이션", level=1)

    add_body(doc, "스킬의 활성화는 두 가지 경로로 이루어진다.")

    add_body_bold(doc, "직접 활성화")
    add_body(
        doc,
        "명시적 슬래시 명령(`/plugin:skill`) 또는 스킬의 메타데이터(frontmatter description)와 "
        "높은 유사도를 보이는 자연어 요청을 런타임이 감지하면, 해당 스킬이 즉시 로드된다. 중간 계층 "
        "없이 사용자의 의도가 특정 스킬에 직결되는 경로이다."
    )

    add_body_bold(doc, "Core 경유 활성화")
    add_body(
        doc,
        "사용자의 요청이 모호하거나 복수 스킬에 걸치는 경우, 런타임은 Core 스킬을 먼저 활성화한다. "
        "Core 스킬은 요청의 의도를 판별하고 적절한 스킬로 라우팅하는 역할에 한정된다. 실제 작업 실행은 "
        "라우팅된 대상 스킬(Orchestrator, Planning 등)이 담당한다."
    )

    add_body(
        doc,
        "스킬 유형에 따라 실행 경로가 분기된다. 위임형 스킬(Core, Planning, Orchestrator)은 "
        "에이전트를 호출하여 작업을 위임하고, 직결형 스킬(Setup, Utility)은 Gateway의 도구를 "
        "직접 사용한다. 이 이중 경로는 Clean Architecture의 관심사 분리와 YAGNI 원칙이 동시에 "
        "적용된 결과이다. 설치나 상태 확인 같은 절차적 작업에 에이전트를 경유하면 불필요한 LLM "
        "호출이 발생하므로, 직결형 스킬이 Gateway에 직접 접근하는 것이 합리적이다."
    )

    add_body(
        doc,
        "런타임은 세션 시작 시 `skills/` 디렉토리를 스캔하여 모든 스킬을 자동 발견한다. 별도의 "
        "등록 절차 없이 디렉토리에 스킬을 배치하면 즉시 사용 가능한 구조이다."
    )

    add_image_if_exists(
        doc,
        "fig_skill_activation.png",
        "스킬 활성화 경로 \u2014 직접 활성화와 Core 경유 활성화의 이중 경로",
    )

    doc.add_paragraph()


def section_06_case_study(doc):
    """섹션 6: 구현 사례."""
    add_heading(doc, "구현 사례 \u2014 DMAP과 abra 플러그인", level=1)

    add_body(
        doc,
        "앞서 설명한 선언형 아키텍처의 실제 구현체가 DMAP(Declarative Multi-Agent Plugin)이다. "
        "DMAP은 코드 없이 Markdown과 YAML만으로 멀티에이전트 시스템을 정의하는 오픈소스 플러그인 "
        "아키텍처 표준이며, DMAP 빌더는 이 표준에 따라 플러그인을 자동 생성하는 도구이다."
    )

    add_body(
        doc,
        "DMAP의 적용 사례로 abra 플러그인이 존재한다. abra는 자연어 한 마디로 Dify 워크플로우 "
        "DSL을 자동 생성하고 프로토타이핑까지 수행하는 플러그인으로, 5개 에이전트"
        "(scenario-analyst, dsl-architect, plan-writer, prototype-runner, agent-developer)가 "
        "각자의 전문 역할을 담당하고, 8개 스킬이 시나리오 분석부터 코드 개발까지의 워크플로우를 "
        "오케스트레이션한다. 전체 구조가 Markdown과 YAML만으로 구성되어 있으며, 에이전트의 역할 "
        "정의, 도구 연결, 실행 순서 등 어떤 부분에서도 프로그래밍 코드가 사용되지 않는다."
    )

    add_body(
        doc,
        "이러한 구현은 선언형 접근이 이론적 가능성이 아닌 실제 동작하는 시스템으로 실현 "
        "가능하다는 것을 보여준다. 프로젝트 전체는 GitHub에 공개되어 있다 "
        "(https://github.com/unicorn-plugins/dmap)."
    )

    add_image_if_exists(
        doc,
        "fig_abra_example.png",
        "abra 플러그인 구조 \u2014 5개 에이전트와 8개 스킬의 선언적 구성",
    )

    doc.add_paragraph()


def section_07_outlook(doc):
    """섹션 7: 전망."""
    add_heading(doc, "전망 \u2014 선언형 에이전트 생태계의 미래", level=1)

    add_body(
        doc,
        "선언형 접근이 에이전트 개발 생태계에 미칠 영향은 세 가지 차원에서 주목된다."
    )

    add_body(
        doc,
        "첫째, 개발자와 비개발자 간 협업 모델의 변화이다. 에이전트 시스템의 설계가 코드에서 "
        "선언으로 전환되면, 도메인 전문가가 에이전트의 역할과 제약을 직접 정의하고 개발자가 "
        "인프라 매핑을 담당하는 분업이 가능해진다. 에이전트 시스템의 설계 참여 범위가 확장된다."
    )

    add_body(
        doc,
        "둘째, 런타임 간 이식성이 에이전트 재사용의 가능성을 열어준다. 하나의 플러그인이 "
        "Claude Code, Codex CLI, Gemini CLI 등 다양한 런타임에서 runtime-mapping.yaml의 "
        "교체만으로 동작할 수 있다면, 에이전트 생태계는 특정 플랫폼에 종속되지 않는 방향으로 "
        "발전할 수 있다."
    )

    add_body(
        doc,
        "셋째, 에이전트 시스템의 표준화 방향이다. Markdown과 YAML이라는 범용 포맷 위에 "
        "에이전트의 역할, 역량, 제약, 핸드오프를 정형화하면, 에이전트 패키지의 검증, 공유, "
        "조합이 체계적으로 이루어질 수 있는 기반이 형성된다."
    )

    add_body(
        doc,
        "코드에서 선언으로의 패러다임 전환은 궁극적으로 \"무엇을 만들 것인가\"와 \"어떻게 "
        "실행할 것인가\"의 관심사를 분리하는 것이다. 이는 소프트웨어 공학이 반복적으로 증명해 온 "
        "원칙 \u2014 추상화 수준의 상승이 시스템의 복잡도를 관리 가능하게 만든다는 것 \u2014 이 "
        "LLM 에이전트 영역에도 유효함을 시사한다."
    )

    add_note(
        doc,
        "DMAP 표준 문서와 샘플은 GitHub에서 확인할 수 있다: "
        "https://github.com/unicorn-plugins/dmap",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # 문서 여백 설정 (2.5cm 사방)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 폰트 설정 (맑은 고딕, 11pt, 행간 1.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 타이틀 (22pt, Bold, 중앙)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("선언형 멀티에이전트 플러그인: 코드 없이 에이전트 시스템을 설계하는 방법")
    run.font.size = Pt(22)
    run.font.name = FONT_NAME
    run.bold = True

    # 서브타이틀 (13pt, 색상 #555555, 중앙)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Clean Architecture 원칙으로 LLM 에이전트 오케스트레이션을 재정의하다")
    run.font.size = Pt(13)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(
        int(COLOR_SUBTITLE[0:2], 16),
        int(COLOR_SUBTITLE[2:4], 16),
        int(COLOR_SUBTITLE[4:6], 16)
    )

    doc.add_paragraph()  # 빈 줄

    # 7개 섹션 순차 호출
    section_01_intro(doc)
    section_02_declarative(doc)
    section_03_architecture(doc)
    section_04_mechanism(doc)
    section_05_orchestration(doc)
    section_06_case_study(doc)
    section_07_outlook(doc)

    # 저장
    doc.save(OUTPUT_PATH)
    print(f"Medium article generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
