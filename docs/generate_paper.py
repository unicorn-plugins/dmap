#!/usr/bin/env python3
"""
DMAP 논문 생성 스크립트
선언형 멀티에이전트 오케스트레이션 논문을 Word(.docx) 형식으로 생성함.

사용법: python docs/generate_paper.py
출력: docs/DMAP-paper.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "DMAP-paper.docx")

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs keys: top, bottom, left, right, insideH, insideV.
    Each value is a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row shaded."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "맑은 고딕"

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_image_if_exists(doc, filename, caption, width_inches=5.5):
    """Add an image with caption if the file exists."""
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        # Skip gracefully
        p = doc.add_paragraph()
        run = p.add_run(f"[그림 누락: {filename}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True


def add_heading(doc, text, level=1):
    """Add heading with appropriate formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
    return h


def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "맑은 고딕"
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_body_bold(doc, text):
    """Add bold body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "맑은 고딕"
    run.bold = True
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = Pt(13)
    # Add gray shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    return p


def add_numbered(doc, number, text, bold_prefix=None):
    """Add a numbered item."""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(f"{number}. {bold_prefix}")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
    else:
        run = p.add_run(f"{number}. {text}")
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
    p.paragraph_format.line_spacing = Pt(15)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


# ===========================================================================
# MAIN: Build the document
# ===========================================================================

def build_document():
    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup: A4, 2.54cm margins
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Page numbers - bottom center
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # -----------------------------------------------------------------------
    # Default style modifications
    # -----------------------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = Pt(15)

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "맑은 고딕"
        if i == 1:
            hs.font.size = Pt(14)
        elif i == 2:
            hs.font.size = Pt(12)
        else:
            hs.font.size = Pt(11)

    # ===================================================================
    # TITLE PAGE
    # ===================================================================
    for _ in range(4):
        doc.add_paragraph()

    # Title (Korean)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("선언형 멀티에이전트 오케스트레이션:")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "맑은 고딕"

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p2.add_run("마크다운과 YAML 기반 Clean Architecture의")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "맑은 고딕"

    title_p3 = doc.add_paragraph()
    title_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p3.add_run("LLM 에이전트 시스템 적용")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "맑은 고딕"

    doc.add_paragraph()

    # Subtitle (English)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(
        "Declarative Multi-Agent Orchestration:\n"
        "Applying Clean Architecture Principles to LLM Agent Systems\n"
        "via Markdown and YAML"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Authors
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = auth_p.add_run("이온디 (Ondi Lee)")
    run.font.size = Pt(11)
    run.font.name = "맑은 고딕"

    affil_p = doc.add_paragraph()
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affil_p.add_run("유니콘주식회사 (Unicorn Inc.)")
    run.font.size = Pt(11)
    run.font.name = "맑은 고딕"

    # Page break after title
    doc.add_page_break()

    # ===================================================================
    # ABSTRACT
    # ===================================================================
    abs_heading = doc.add_paragraph()
    abs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = abs_heading.add_run("초록 (Abstract)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "맑은 고딕"

    abstract_text = (
        "LLM(대규모 언어 모델) 기반 멀티에이전트 시스템은 2023년 이후 급격한 성장을 보이며, "
        "LangChain, CrewAI, AutoGen, MetaGPT 등 다양한 프레임워크가 등장하였다. "
        "그러나 이들 프레임워크는 Python/TypeScript SDK에 대한 코드 종속성, "
        "특정 프레임워크에 대한 런타임 결합, 오케스트레이션과 실행의 역할 혼재, "
        "도구 및 모델의 추상화 부족 등 공통적인 구조적 한계를 가진다. "
        "본 연구는 소프트웨어 공학에서 검증된 Clean Architecture의 3대 원칙"
        "(Loosely Coupling, High Cohesion, Dependency Inversion)을 "
        "코드가 아닌 선언형 명세(Markdown + YAML)로 LLM 에이전트 시스템에 적용하는 "
        "DMAP(Declarative Multi-Agent Plugin) 아키텍처를 제안한다. "
        "본 연구의 핵심 기여는 다음 네 가지이다: "
        "(1) 마크다운과 YAML만으로 멀티에이전트 시스템을 정의하는 선언형 표준 제안, "
        "(2) Clean Architecture 원칙의 AI 에이전트 시스템 체계적 이식, "
        "(3) 4-Tier 모델과 Gateway 매핑을 통한 런타임 중립적 추상 계층 설계, "
        "(4) 실증적 검증. "
        "OMC(Oh-My-ClaudeCode) 플러그인(39개 스킬, 35개 에이전트)과 "
        "Abra 플러그인(비즈니스 도메인 AI Agent 개발 워크플로우)의 실제 운용 사례를 통해 "
        "표준의 실효성을 검증하였다. "
        "본 아키텍처는 비개발자의 에이전트 시스템 참여를 가능하게 하고, "
        "런타임 간 이식성을 확보하며, 도메인에 무관한 범용 적용 가능성을 제시한다."
    )

    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.left_indent = Cm(1.5)
    abs_p.paragraph_format.right_indent = Cm(1.5)
    run = abs_p.add_run(abstract_text)
    run.font.size = Pt(10)
    run.font.name = "맑은 고딕"
    abs_p.paragraph_format.line_spacing = Pt(14)

    doc.add_paragraph()

    # Keywords
    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.left_indent = Cm(1.5)
    run = kw_p.add_run("키워드: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "맑은 고딕"
    run = kw_p.add_run(
        "멀티에이전트 시스템, 선언형 아키텍처, Clean Architecture, "
        "LLM 오케스트레이션, 런타임 중립성, 플러그인 표준"
    )
    run.font.size = Pt(10)
    run.font.name = "맑은 고딕"

    doc.add_page_break()

    # ===================================================================
    # SECTION 1: 서론
    # ===================================================================
    add_heading(doc, "1. 서론 (Introduction)", level=1)

    add_heading(doc, "1.1 연구 배경", level=2)

    add_body(doc,
        "대규모 언어 모델(Large Language Model, LLM)은 2023년을 기점으로 "
        "단순한 텍스트 생성 도구를 넘어 자율적 에이전트(Autonomous Agent) 시스템의 "
        "핵심 엔진으로 진화하였다. GPT-4, Claude, Gemini 등 최신 LLM의 추론 능력 향상에 힘입어, "
        "복수의 LLM 에이전트가 협업하여 복잡한 작업을 수행하는 "
        "멀티에이전트 시스템(Multi-Agent System)이 급격히 성장하고 있다."
    )

    add_body(doc,
        "이 흐름을 주도하는 대표적 프레임워크로는 LangChain/LangGraph [1], "
        "CrewAI [2], AutoGen [3], MetaGPT [4], ChatDev [5] 등이 있다. "
        "이들은 에이전트 간 협업, 역할 분담, 워크플로우 오케스트레이션이라는 핵심 문제를 "
        "각각의 방식으로 해결하고 있으나, 공통적인 구조적 한계를 내포한다."
    )

    add_heading(doc, "1.2 기존 프레임워크의 한계", level=2)

    add_body(doc,
        "기존 프레임워크의 공통적 한계는 다음 다섯 가지로 요약된다."
    )

    make_table(doc,
        headers=["문제", "설명", "결과"],
        rows=[
            ["코드 종속성", "Python/TypeScript SDK로만 에이전트 정의 가능",
             "비개발자 진입 불가, 도메인 전문가 배제"],
            ["런타임 결합", "특정 프레임워크에 강결합",
             "다른 런타임으로 이식 불가"],
            ["역할 혼재", "에이전트가 오케스트레이션과 실행을 동시 수행",
             "관심사 분리 부재, 유지보수 어려움"],
            ["추상화 부족", "도구, 모델이 코드에 하드코딩",
             "환경 변경 시 코드 수정 필요"],
            ["도메인 한정", "대부분 코드 생성/데이터 분석에 특화",
             "교육, 문서화, 비즈니스 워크플로우 적용 어려움"],
        ],
        col_widths=[3, 6, 6]
    )

    add_body(doc, "")  # spacer

    add_body(doc,
        "특히 모든 기존 프레임워크가 Python 또는 TypeScript SDK를 필수적으로 요구한다는 점은, "
        "도메인 전문가(교육자, 비즈니스 분석가 등)가 에이전트 시스템의 설계에 직접 참여하지 못하는 "
        "근본적 장벽으로 작용한다. 또한 특정 프레임워크에 결합된 에이전트 정의는 "
        "다른 런타임 환경으로 이식이 불가능하여, 기술 종속(vendor lock-in) 문제를 야기한다."
    )

    add_heading(doc, "1.3 연구 질문 및 목적", level=2)

    add_body(doc,
        "본 연구는 다음과 같은 핵심 연구 질문을 제기한다:"
    )

    q_p = doc.add_paragraph()
    q_p.paragraph_format.left_indent = Cm(1.5)
    q_p.paragraph_format.right_indent = Cm(1.5)
    run = q_p.add_run(
        '"소프트웨어 공학의 검증된 원칙(Clean Architecture)을 '
        '코드 작성 없이 AI 에이전트 시스템에 적용할 수 있는가?"'
    )
    run.font.size = Pt(10.5)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "이 질문에 답하기 위해, 본 연구는 DMAP(Declarative Multi-Agent Plugin) 아키텍처를 "
        "제안한다. DMAP은 마크다운(프롬프트)과 YAML(설정)만으로 멀티에이전트 시스템을 "
        "선언적으로 정의하는 범용 플러그인 표준이다."
    )

    add_heading(doc, "1.4 핵심 기여", level=2)

    add_body(doc, "본 연구의 핵심 기여는 다음 네 가지이다:")

    add_numbered(doc, 1, " 마크다운과 YAML만으로 멀티에이전트 시스템을 정의하는 "
        "선언형 에이전트 아키텍처 표준을 제안한다. 코드 작성 없이 에이전트의 역할, 역량, "
        "제약, 핸드오프를 선언적으로 명세할 수 있다.",
        bold_prefix="선언형 에이전트 아키텍처 표준 제안.")
    add_numbered(doc, 2, " Loosely Coupling, High Cohesion, Dependency Inversion 원칙을 "
        "LLM 에이전트 오케스트레이션에 체계적으로 적용한다.",
        bold_prefix="Clean Architecture의 AI 에이전트 시스템 이식.")
    add_numbered(doc, 3, " 4-Tier 모델, Gateway 매핑 등 추상 선언과 구체 매핑을 분리하여 "
        "Claude Code, Codex CLI, Gemini CLI 등 어떤 런타임에서도 동일한 플러그인이 "
        "동작하는 이식성을 확보한다.",
        bold_prefix="런타임 중립적 추상 계층 설계.")
    add_numbered(doc, 4, " OMC 플러그인(39 Skills, 35 Agents)과 "
        "Abra 플러그인(비즈니스 도메인)의 실제 운용 사례를 통해 표준의 실효성을 검증한다.",
        bold_prefix="실증적 검증.")

    add_body(doc,
        "본 논문의 나머지 구성은 다음과 같다. 2장에서 관련 연구를 분석하고, "
        "3장에서 DMAP 아키텍처를, 4장에서 설계 원칙을, 5장에서 구현 세부사항을 기술한다. "
        "6장에서 사례 연구, 7장에서 평가, 8장에서 논의, 9장에서 결론을 제시한다."
    )

    # ===================================================================
    # SECTION 2: 관련 연구
    # ===================================================================
    add_heading(doc, "2. 관련 연구 (Related Work)", level=1)

    add_heading(doc, "2.1 LLM 기반 멀티에이전트 프레임워크", level=2)

    add_body(doc,
        "LLM 기반 멀티에이전트 시스템 분야에서 주요 프레임워크들의 접근 방식과 "
        "특성을 분석한다."
    )

    add_body_bold(doc, "LangChain / LangGraph [1]")
    add_body(doc,
        "LangChain은 LLM 애플리케이션 개발의 사실상 표준 프레임워크로, "
        "도구 체인(tool chain) 방식으로 에이전트를 구성한다. "
        "LangGraph는 이를 확장하여 그래프 기반 워크플로우를 지원하며, "
        "상태 기계(state machine) 패턴으로 복잡한 에이전트 간 상호작용을 모델링한다. "
        "그러나 모든 에이전트 정의와 워크플로우가 Python 코드로 작성되어야 하며, "
        "LangChain SDK에 강하게 결합되어 있다."
    )

    add_body_bold(doc, "CrewAI [2]")
    add_body(doc,
        "CrewAI는 역할 기반(role-based) 에이전트 팀 구성을 특징으로 한다. "
        "'크루(crew)' 단위로 에이전트를 조직하고, 역할(role), 목표(goal), "
        "백스토리(backstory)를 Python 코드로 정의한다. 순차적(sequential) 또는 "
        "계층적(hierarchical) 프로세스를 지원하나, CrewAI SDK에 종속적이며 "
        "위임(delegation) 키워드를 제외하면 핸드오프 메커니즘이 부재하다."
    )

    add_body_bold(doc, "AutoGen (Microsoft) [3]")
    add_body(doc,
        "AutoGen은 대화 프로토콜(conversational protocol) 기반의 멀티에이전트 프레임워크로, "
        "에이전트 간 메시지 교환을 통해 협업을 구현한다. "
        "유연한 대화 패턴을 제공하나, 모든 에이전트와 대화 흐름이 "
        "Python 코드로 정의되어야 하며, AutoGen SDK에 강결합된다."
    )

    add_body_bold(doc, "MetaGPT [4]")
    add_body(doc,
        "MetaGPT는 SOP(Standard Operating Procedure) 기반 접근으로, "
        "소프트웨어 개발 조직의 역할(PM, Architect, Developer 등)을 에이전트에 매핑한다. "
        "구조화된 산출물 교환을 통해 높은 품질의 출력을 생성하나, "
        "코드 생성 도메인에 특화되어 있어 범용성이 제한적이다."
    )

    add_body_bold(doc, "ChatDev [5]")
    add_body(doc,
        "ChatDev는 소프트웨어 회사를 시뮬레이션하는 접근으로, "
        "CEO, CTO, 프로그래머, 테스터 등의 역할을 LLM 에이전트에 부여한다. "
        "채팅 체인(chat chain) 방식으로 소프트웨어 개발 프로세스를 자동화하나, "
        "소프트웨어 개발이라는 특정 도메인에 한정된다."
    )

    add_heading(doc, "2.2 비교 분석", level=2)

    add_body(doc, "주요 프레임워크와 본 연구(DMAP)의 비교는 Table 1과 같다.")

    # Table 1: Framework comparison
    make_table(doc,
        headers=["비교 항목", "LangChain/\nLangGraph", "CrewAI", "AutoGen",
                 "MetaGPT", "본 연구\n(DMAP)"],
        rows=[
            ["에이전트 정의 방식", "Python 코드", "Python 코드", "Python 코드",
             "Python 코드", "Markdown\n+ YAML"],
            ["오케스트레이션", "그래프 코드", "순차/계층\n코드", "대화\n프로토콜",
             "SOP 코드", "스킬\n프롬프트"],
            ["런타임 종속성", "LangChain\nSDK", "CrewAI\nSDK", "AutoGen\nSDK",
             "MetaGPT\nSDK", "런타임\n중립"],
            ["아키텍처 원칙", "없음\n(도구 체인)", "역할 기반", "대화 기반",
             "SOP 기반", "Clean\nArchitecture"],
            ["도구 추상화", "Tool 클래스", "Tool\n데코레이터", "Function\ncall",
             "Tool 클래스", "추상 선언\n+ Gateway"],
            ["티어 관리", "없음", "없음", "없음", "없음", "4-Tier +\nruntime-\nmapping"],
            ["핸드오프", "없음", "위임\n키워드", "없음", "없음", "agentcard\n선언"],
            ["도메인 범용성", "중간", "중간", "중간", "코드 특화", "완전 범용"],
            ["비개발자 접근성", "불가", "불가", "불가", "불가", "가능"],
            ["이식성", "낮음", "낮음", "낮음", "낮음", "높음"],
        ],
        col_widths=[2.5, 2.3, 2.0, 2.0, 2.0, 2.3]
    )

    cap_t1 = doc.add_paragraph()
    cap_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t1.add_run("Table 1. 기존 멀티에이전트 프레임워크와의 비교 요약")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_heading(doc, "2.3 패러다임 전환", level=2)

    add_body(doc,
        "기존 프레임워크와 본 연구의 근본적 차이는 "
        "\"코드로 구현(imperative)\"에서 \"선언으로 명세(declarative)\"로의 패러다임 전환에 있다. "
        "기존 접근에서는 개발자가 Python 코드를 작성하여 프레임워크가 에이전트를 실행하는 반면, "
        "본 연구에서는 누구나 Markdown/YAML을 작성하면 어떤 런타임이든 에이전트를 실행할 수 있다."
    )

    make_table(doc,
        headers=["차원", "기존 (Imperative)", "본 연구 (Declarative)"],
        rows=[
            ["에이전트 정의", '"어떻게 동작하는가" 코딩', '"무엇을 할 수 있는가" 선언'],
            ["도구 연결", "코드에서 도구 객체 생성", "tools.yaml에서 추상 인터페이스 선언"],
            ["역할 제약", "if문으로 분기", "forbidden_actions 블랙리스트"],
            ["모델 선택", "코드에서 모델명 하드코딩", "tier 선언 → runtime-mapping.yaml 해석"],
            ["워크플로우", "함수 호출 체인", "프롬프트에 자연어로 기술"],
        ],
        col_widths=[3, 5.5, 5.5]
    )

    cap_t2 = doc.add_paragraph()
    cap_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t2.add_run("Table 2. 명령형 vs 선언형 패러다임 비교")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "이러한 패러다임 전환은 에이전트 시스템의 접근성, 이식성, 유지보수성에 "
        "근본적인 변화를 가져온다. 코드 능력이 없는 도메인 전문가도 "
        "에이전트의 역할과 역량을 직접 명세할 수 있으며, "
        "동일한 선언 파일이 다양한 런타임 환경에서 재사용될 수 있다."
    )

    # ===================================================================
    # SECTION 3: 아키텍처
    # ===================================================================
    add_heading(doc, "3. 아키텍처 (Architecture)", level=1)

    add_heading(doc, "3.1 Clean Architecture 원칙의 적용", level=2)

    add_body(doc,
        "DMAP 아키텍처는 Robert C. Martin이 제안한 Clean Architecture [6]의 "
        "3대 원칙을 LLM 에이전트 오케스트레이션에 체계적으로 적용한다. "
        "Table 3은 각 원칙이 전통적 소프트웨어와 본 연구에서 어떻게 대응되는지 보여준다."
    )

    make_table(doc,
        headers=["원칙", "전통적 소프트웨어", "본 연구의 적용", "구현 메커니즘"],
        rows=[
            ["Loosely\nCoupling",
             "인터페이스를 통한\n의존성 분리",
             "추상 선언(agentcard.yaml\n/ tools.yaml) ↔\n구체 구현(runtime-\nmapping.yaml)",
             "에이전트는 file_read 선언 →\nGateway가 Read 도구로 매핑"],
            ["High\nCohesion",
             "클래스가 단일 책임에\n집중",
             "Skills=라우팅/오케스트레이션\nAgents=자율 실행\nGateway=도구 매핑",
             "각 컴포넌트가\n자기 역할에만 집중"],
            ["Dependency\nInversion",
             "상위 모듈이\n추상에 의존",
             "상위(Skills/Agents)가\n추상에 의존,\n하위(Gateway/Runtime)가\n구체 제공",
             "tier: HIGH →\nruntime-mapping.yaml이\n모델로 해석"],
        ],
        col_widths=[2.5, 3.5, 4, 4]
    )

    cap_t3 = doc.add_paragraph()
    cap_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t3.add_run("Table 3. Clean Architecture 3대 원칙의 적용")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_heading(doc, "3.2 5-Layer 아키텍처", level=2)

    add_body(doc,
        "DMAP은 5개의 수직 계층(Layer)과 1개의 횡단 관심사(Cross-cutting Concern)로 구성된다. "
        "Figure 1은 전체 아키텍처의 계층 구조를 보여준다."
    )

    add_image_if_exists(doc, "fig_architecture.png",
                        "Figure 1. DMAP 5-Layer 아키텍처")

    add_body(doc,
        "최상위 계층인 Input은 사용자의 요청이 진입하는 지점이다. "
        "Controller + UseCase 계층의 Skills는 사용자 요청을 수신하여 적절한 "
        "에이전트에게 작업을 위임하거나, 직접 Gateway 도구를 사용하는 라우팅 및 "
        "오케스트레이션을 수행한다. Service 계층의 Agents는 위임받은 작업을 "
        "자율적으로 수행하는 전문가 역할이며, Gateway 계층은 추상 선언을 "
        "구체적인 런타임 도구로 매핑하는 인프라 역할을 담당한다. "
        "최하위 Runtime 계층은 실제 실행 환경(Claude Code, Codex CLI, Gemini CLI 등)이다."
    )

    add_body_bold(doc, "스킬 실행 경로")

    add_body(doc,
        "DMAP은 두 가지 스킬 실행 경로를 정의한다:"
    )

    add_bullet(doc,
        " LLM 추론이 필요한 작업에 대해 Agent에 위임하는 경로이다. "
        "Input → Skills(Controller) → Agents(Service) → Gateway → Runtime의 "
        "전체 계층을 경유한다. Core, Planning, Orchestrator 유형의 스킬이 이 경로를 사용한다.",
        bold_prefix="위임형 경로(Delegated Path):")
    add_bullet(doc,
        " 절차적이고 결정론적인 작업에 대해 Agent 계층을 생략하고 "
        "Gateway 도구를 직접 사용하는 경로이다. "
        "Input → Skills(Controller) → Gateway → Runtime으로 Service 계층을 우회한다. "
        "Setup, Utility 유형의 스킬이 이 경로를 사용하며, "
        "이는 YAGNI(You Aren't Gonna Need It) 원칙의 적용으로, "
        "불필요한 LLM 호출을 방지한다.",
        bold_prefix="직결형 경로(Direct Path):")

    add_body(doc,
        "추가로, Hooks는 Cross-cutting Concern(AOP, Aspect-Oriented Programming)으로서 "
        "모든 계층의 이벤트를 횡단적으로 가로챈다. "
        "단, Hooks는 오케스트레이션 플러그인(예: OMC)에서만 사용이 허용되며, "
        "일반 플러그인에서는 사용이 금지된다."
    )

    add_heading(doc, "3.3 4-Tier 에이전트 모델", level=2)

    add_body(doc,
        "DMAP은 동일 역할의 에이전트를 비용-역량 트레이드오프에 따라 4개의 티어로 분리하는 "
        "범용 원칙을 제안한다. 티어(tier)는 LLM 모델 등급을 결정하는 추상 선언이며, "
        "Gateway의 runtime-mapping.yaml이 실제 모델명으로 매핑한다. "
        "Figure 2는 4-Tier 모델의 구조를 보여준다."
    )

    add_image_if_exists(doc, "fig_tier_model.png",
                        "Figure 2. 4-Tier 에이전트 모델")

    make_table(doc,
        headers=["티어", "특성", "LLM 예시", "적합 작업", "에스컬레이션"],
        rows=[
            ["LOW", "빠르고 저비용", "Haiku", "단건 조회, 간단한 수정",
             "복잡도 초과 시\n상위 티어로 보고"],
            ["MEDIUM", "균형", "Sonnet", "기능 구현, 일반 분석", "—"],
            ["HIGH", "최고 역량, 고비용", "Opus", "복잡한 의사결정,\n심층 분석", "—"],
            ["HEAVY", "최고 역량 +\n대규모 예산", "Opus\n(대규모 토큰·시간)",
             "장시간 추론,\n대규모 멀티파일 작업", "—"],
        ],
        col_widths=[2, 2.5, 2.5, 3.5, 3.5]
    )

    cap_t4 = doc.add_paragraph()
    cap_t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t4.add_run("Table 4. 4-Tier 에이전트 모델")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "핵심 메커니즘으로 에스컬레이션(escalation)이 있다. "
        "LOW 티어 에이전트가 자신의 역량 한계를 인식하면, "
        "상위 티어로 작업을 보고하는 메커니즘이다. "
        "이는 IT 서비스 관리의 L1→L2→L3 지원 체계와 동일한 원리이다. "
        "또한 티어 변형 에이전트는 기본 에이전트의 설정을 상속하고, "
        "오버라이드가 필요한 부분만 기술하여 선언 중복을 최소화한다."
    )

    add_heading(doc, "3.4 3계층 활성화 구조", level=2)

    add_body(doc,
        "DMAP은 스킬 활성화의 순환 의존 문제를 해결하기 위해 "
        "3계층 활성화 구조를 설계하였다. Figure 3은 이 구조를 보여준다."
    )

    add_image_if_exists(doc, "fig_activation.png",
                        "Figure 3. 3계층 활성화 구조")

    add_body(doc,
        "활성화 조건이 스킬 내부에만 존재하면, 스킬을 로드하기 위해 "
        "조건을 알아야 하는 순환 의존(circular dependency) 문제가 발생한다. "
        "DMAP은 이를 3계층으로 분리하여 해결한다:"
    )

    add_numbered(doc, 1, " 항상 로드되는 런타임 상주 파일(예: CLAUDE.md)에 "
        "라우팅 테이블을 배치하여, \"이 요청은 어떤 플러그인의 어떤 스킬인가?\"를 판단한다.",
        bold_prefix="런타임 상주 파일(Routing Table).")
    add_numbered(doc, 2, " 조건 매칭 시에만 로드되는 핵심 스킬이 "
        "runtime-mapping.yaml을 참조하여 에이전트를 스폰한다.",
        bold_prefix="핵심 스킬(Core Skill).")
    add_numbered(doc, 3, " 위임받아 자율적으로 작업을 수행하고 결과를 반환한다.",
        bold_prefix="에이전트(Agent).")

    add_body(doc,
        "Setup 스킬은 플러그인 설치 시 런타임 상주 파일에 Core 스킬의 활성화 조건을 "
        "라우팅 테이블로 등록하는 역할을 수행한다. 이를 통해 플러그인의 자기 등록(self-registration) "
        "메커니즘이 구현된다."
    )

    # ===================================================================
    # SECTION 4: 설계 원칙
    # ===================================================================
    add_heading(doc, "4. 설계 원칙 (Design Principles)", level=1)

    add_body(doc,
        "DMAP 아키텍처는 11가지 설계 포인트를 정의하며, "
        "각 포인트는 소프트웨어 공학의 검증된 원칙에 대응된다. "
        "Table 5는 전체 설계 포인트의 요약이다."
    )

    make_table(doc,
        headers=["#", "설계 포인트", "설명", "소프트웨어 공학\n대응 개념"],
        rows=[
            ["1", "런타임 중립성",
             "tier: HIGH라는 추상 선언으로\nLLM, 클라우드, 인력 등\n도메인 무관 적용",
             "Dependency\nInversion Principle"],
            ["2", "3계층 활성화",
             "런타임 상주 파일(라우팅) →\n핵심 스킬(오케스트레이션) →\n에이전트(실행)의 순환 의존 해결",
             "Layered\nArchitecture"],
            ["3", "프롬프트 깊이\n차등화",
             "라우팅/분기는 상세히,\n에이전트 위임은 간결하게\n(WHAT만, HOW는 에이전트 자율)",
             "Interface\nSegregation"],
            ["4", "위임 표기법",
             "Agent 위임 5항목\n(TASK/EXPECTED OUTCOME/\nMUST DO/MUST NOT DO/CONTEXT)\nSkill 위임 3항목\n(INTENT/ARGS/RETURN)",
             "Command Pattern"],
            ["5", "에스컬레이션",
             "LOW 티어가 자기 한계 인식 →\n상위 티어 위임",
             "L1→L2→L3\n지원 체계"],
            ["6", "install.yaml ↔\nsetup 스킬 분리",
             "데이터(WHAT) / 지시(HOW) /\n실행(DO) 3단 분리",
             "CQRS"],
            ["7", "핸드오프 선언",
             "agentcard.yaml에\ntarget + when + reason으로\n역할 경계 명시",
             "Service Contract"],
            ["8", "액션 카테고리\n추상화",
             "file_write 선언 →\nruntime-mapping.yaml이\nWrite, Edit 매핑",
             "Adapter Pattern"],
            ["9", "에이전트 패키지\n경계 원칙",
             "AGENT.md(WHY+HOW)와\nagentcard.yaml(WHO+WHAT\n+WHEN)에 동일 정보 중복 금지",
             "Separation of\nConcerns"],
            ["10", "직결형 경로\n(YAGNI)",
             "Setup/Utility 스킬은\nAgent 계층 불필요 →\nGateway 직접 접근",
             "YAGNI (XP)"],
            ["11", "프롬프트\n조립 순서",
             "공통 정적 → 에이전트별 정적\n→ 동적 순서로\nprefix 캐시 적중률 극대화",
             "Cache\nOptimization"],
        ],
        col_widths=[0.8, 2.5, 5, 3]
    )

    cap_t5 = doc.add_paragraph()
    cap_t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t5.add_run("Table 5. DMAP 11가지 설계 포인트")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_heading(doc, "4.1 런타임 중립성과 Dependency Inversion", level=2)

    add_body(doc,
        "DMAP의 가장 핵심적인 설계 원칙은 런타임 중립성이다. "
        "에이전트 패키지(agentcard.yaml, tools.yaml)는 'tier: HIGH', "
        "'code_search' 등 추상 선언만을 포함하며, 이를 해석하여 실제 모델명(claude-opus-4-6)이나 "
        "도구(lsp_workspace_symbols)로 매핑하는 것은 Gateway의 runtime-mapping.yaml이 담당한다. "
        "이는 Dependency Inversion Principle의 직접적 적용으로, "
        "상위 모듈(Skills, Agents)이 추상에 의존하고, "
        "하위 모듈(Gateway, Runtime)이 구체 구현을 제공하는 구조이다."
    )

    add_body(doc,
        "이 원칙의 실질적 효과는, 동일한 에이전트 패키지가 Claude Code 환경에서는 "
        "claude-opus-4-6 모델과 lsp_workspace_symbols 도구로 실행되고, "
        "Codex CLI 환경에서는 o3 모델과 해당 런타임의 검색 도구로 실행될 수 있다는 것이다. "
        "runtime-mapping.yaml만 교체하면 나머지 모든 선언 파일은 변경 없이 재사용된다."
    )

    add_heading(doc, "4.2 프롬프트 깊이 차등화와 위임 표기법", level=2)

    add_body(doc,
        "DMAP에서 스킬은 두 가지 수준의 프롬프트 깊이를 사용한다. "
        "라우팅과 분기 로직은 상세하게 기술하여 정확한 판단을 유도하고, "
        "에이전트 위임 시에는 WHAT(무엇을 해야 하는가)만 전달하고 "
        "HOW(어떻게 수행하는가)는 에이전트의 자율에 맡긴다. "
        "이는 Interface Segregation Principle의 응용으로, "
        "각 계층에 필요한 최소한의 정보만 전달하는 원칙이다."
    )

    add_body(doc,
        "위임 표기법은 두 가지로 구분된다. Agent 위임은 TASK(작업 지시), "
        "EXPECTED OUTCOME(기대 결과), MUST DO(필수 행동), MUST NOT DO(금지 행동), "
        "CONTEXT(맥락 정보)의 5항목으로 구성된다. "
        "Skill 위임은 INTENT(의도), ARGS(인자), RETURN(반환)의 3항목으로 구성되며, "
        "이는 Command Pattern의 적용이다."
    )

    add_heading(doc, "4.3 경계 원칙과 관심사 분리", level=2)

    add_body(doc,
        "에이전트 패키지의 경계 원칙은 AGENT.md와 agentcard.yaml에 "
        "동일 정보를 중복 기술하지 않는 것이다. "
        "AGENT.md는 LLM이 읽는 프롬프트로서 WHY(목표)와 HOW(워크플로우)를 담당하고, "
        "agentcard.yaml은 런타임이 읽는 메타데이터로서 WHO(정체성), "
        "WHAT(역량, 제약), WHEN(핸드오프)을 담당한다. "
        "이 분리는 Separation of Concerns 원칙의 직접적 적용이다."
    )

    add_body(doc,
        "더불어 install.yaml과 setup 스킬의 분리는 데이터(WHAT: 무엇을 설치하는가), "
        "지시(HOW: 어떻게 설치하는가), 실행(DO: 실제 설치 수행)의 3단 분리를 구현하며, "
        "이는 CQRS(Command Query Responsibility Segregation) 원리와 유사하다."
    )

    # ===================================================================
    # SECTION 5: 구현
    # ===================================================================
    add_heading(doc, "5. 구현 (Implementation)", level=1)

    add_heading(doc, "5.1 스킬 유형 체계", level=2)

    add_body(doc,
        "DMAP은 5가지 스킬 유형을 정의하며, 각 유형은 실행 경로(위임형/직결형)와 "
        "역할이 명확히 구분된다."
    )

    make_table(doc,
        headers=["유형", "실행 경로", "역할", "필수 여부", "플러그인당 수"],
        rows=[
            ["Core\n(핵심)", "위임형", "사용자 요청 라우팅,\n에이전트 위임", "필수", "1개"],
            ["Setup\n(설정)", "직결형", "플러그인 설치,\n라우팅 테이블 등록", "필수", "1개 이상"],
            ["Planning\n(계획)", "위임형", "전략 계획 수립,\n요구사항 분석", "선택", "제한 없음"],
            ["Orchestrator\n(지휘자)", "위임형", "멀티에이전트\n워크플로우 조율", "선택", "제한 없음"],
            ["Utility\n(유틸리티)", "직결형", "보조 기능,\n도구 직접 사용", "선택", "제한 없음"],
        ],
        col_widths=[2, 2, 3.5, 2, 2.5]
    )

    cap_t6 = doc.add_paragraph()
    cap_t6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t6.add_run("Table 6. 5가지 스킬 유형")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "위임형 스킬(Core, Planning, Orchestrator)은 LLM 추론이 필요한 작업을 "
        "에이전트에게 위임하며, 자신은 라우팅과 오케스트레이션에만 집중한다. "
        "직결형 스킬(Setup, Utility)은 절차적이고 결정론적인 작업을 "
        "Agent 계층 없이 Gateway 도구를 직접 사용하여 수행한다. "
        "모든 플러그인은 반드시 하나의 Core 스킬과 하나 이상의 Setup 스킬을 포함해야 한다."
    )

    add_heading(doc, "5.2 에이전트 패키지 구조", level=2)

    add_body(doc,
        "에이전트는 독립된 디렉토리 패키지로 구성되며, "
        "3개의 핵심 파일이 명확한 역할 분리를 이룬다. "
        "Figure 4는 에이전트 패키지의 구조를 보여준다."
    )

    add_image_if_exists(doc, "fig_agent_package.png",
                        "Figure 4. 에이전트 패키지 구조")

    make_table(doc,
        headers=["파일", "독자", "관점", "내용"],
        rows=[
            ["AGENT.md", "LLM\n(프롬프트 주입)", "WHY + HOW",
             "목표, 워크플로우, 출력 형식, 검증"],
            ["agentcard.yaml", "런타임\n(기계 판독)", "WHO + WHAT\n+ WHEN",
             "정체성(is/is_not), 역량, 제약\n(forbidden_actions), 핸드오프\n(target+when+reason), 에스컬레이션"],
            ["tools.yaml", "런타임\n(매칭 참조)", "WHAT\n(도구)",
             "추상 도구 인터페이스:\nname, description, input, output"],
        ],
        col_widths=[2.5, 2.5, 2.5, 6]
    )

    cap_t7 = doc.add_paragraph()
    cap_t7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t7.add_run("Table 7. 에이전트 패키지 3파일 역할 분리")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "AGENT.md는 Frontmatter(메타데이터) → 목표(## 목표) → 참조(## 참조) → "
        "워크플로우(## 워크플로우) → 출력 형식(## 출력 형식) → 검증(## 검증)의 "
        "표준 섹션 구조를 따른다. AGENT.md에서 도구를 참조할 때는 "
        "{tool:name} 표기법으로 tools.yaml에 정의된 추상 도구를 참조하며, "
        "모델명이나 구체 도구명을 하드코딩하지 않는다."
    )

    add_body(doc,
        "agentcard.yaml은 name, version, tier, capabilities(role, identity, restrictions), "
        "handoff(target, when, reason), escalation 등의 필드를 기계 판독 가능한 YAML 형식으로 선언한다. "
        "identity는 is(긍정적 정체성)와 is_not(부정적 정체성)으로 에이전트의 역할 범위를 "
        "명확히 하며, restrictions의 forbidden_actions는 블랙리스트 방식으로 "
        "에이전트가 수행해서는 안 되는 행동을 선언한다."
    )

    add_heading(doc, "5.3 위임 표기법", level=2)

    add_body(doc,
        "DMAP은 두 가지 위임 표기법을 정의한다."
    )

    add_body_bold(doc, "Agent 위임 (5항목)")

    make_table(doc,
        headers=["항목", "역할", "예시"],
        rows=[
            ["TASK", "수행할 작업 지시", "\"auth 모듈의 JWT 검증 로직을 구현하라\""],
            ["EXPECTED\nOUTCOME", "기대하는 결과물", "\"auth.ts에 validateToken 함수 추가\""],
            ["MUST DO", "반드시 수행할 행동", "\"단위 테스트 포함, 에러 처리 구현\""],
            ["MUST NOT DO", "금지 행동", "\"기존 API 시그니처 변경 금지\""],
            ["CONTEXT", "참고할 맥락 정보", "\"현재 Express.js 기반, TypeScript 사용\""],
        ],
        col_widths=[2.5, 3.5, 8]
    )

    add_body(doc, "")  # spacer

    add_body_bold(doc, "Skill 위임 (3항목)")

    make_table(doc,
        headers=["항목", "역할", "예시"],
        rows=[
            ["INTENT", "위임 의도", "\"빌드 오류를 수정하라\""],
            ["ARGS", "전달 인자", "\"대상 디렉토리: src/\""],
            ["RETURN", "기대 반환", "\"수정된 파일 목록과 변경 내역\""],
        ],
        col_widths=[2.5, 3.5, 8]
    )

    add_heading(doc, "5.4 Gateway 구현", level=2)

    add_body(doc,
        "Gateway는 추상 선언과 구체 구현 사이의 매핑을 담당하며, "
        "install.yaml과 runtime-mapping.yaml의 2파일로 구성된다. "
        "Figure 5는 Gateway의 추상-구체 분리 구조를 보여준다."
    )

    add_image_if_exists(doc, "fig_gateway.png",
                        "Figure 5. Gateway 추상-구체 분리 구조")

    add_body(doc,
        "install.yaml은 플러그인이 필요로 하는 외부 도구(MCP 서버, LSP 서버, 커스텀 도구)의 "
        "설치 매니페스트이다. \"무엇을(WHAT) 설치하는가\"만을 선언하며, "
        "\"어떻게(HOW) 설치하는가\"는 Setup 스킬이 담당한다."
    )

    add_body(doc,
        "runtime-mapping.yaml은 3개의 매핑 영역으로 구성된다:"
    )

    make_table(doc,
        headers=["영역", "역할", "입력", "출력"],
        rows=[
            ["tier_mapping", "티어 → 모델 매핑",
             "agentcard.yaml의 tier", "실제 LLM 모델명"],
            ["tool_mapping", "추상 도구 → 실제 도구",
             "tools.yaml의 도구 선언", "lsp/mcp/custom 도구"],
            ["action_mapping", "금지 액션 → 실제 도구",
             "agentcard.yaml의\nforbidden_actions", "제외할 실제 도구 목록"],
        ],
        col_widths=[2.5, 3.5, 4, 4]
    )

    cap_t8 = doc.add_paragraph()
    cap_t8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t8.add_run("Table 8. runtime-mapping.yaml 3영역")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "tool_mapping에는 lsp, mcp, custom 타입만 매핑하며, "
        "builtin 도구(Read, Write, Bash 등)는 런타임이 내장 처리하므로 매핑에서 생략한다. "
        "예산(budget) 관리는 별도 파일 없이 런타임이 자체 관리한다."
    )

    add_heading(doc, "5.5 프롬프트 조립 3단계", level=2)

    add_body(doc,
        "위임형 스킬이 에이전트를 스폰할 때, 프롬프트는 다음 3단계로 조립된다. "
        "이 순서는 런타임의 prefix 캐시 적중률을 극대화하기 위해 설계되었다."
    )

    make_table(doc,
        headers=["순서", "단계", "내용", "캐시 특성"],
        rows=[
            ["1", "공통 정적", "runtime-mapping.yaml\n(tier, tool, action 매핑)",
             "모든 에이전트 공통\n→ 높은 캐시 적중률"],
            ["2", "에이전트별\n정적", "AGENT.md + agentcard.yaml\n+ tools.yaml (3파일)",
             "에이전트별 고정\n→ 동일 에이전트 호출 시\n캐시 적중"],
            ["3", "동적", "작업 지시\n(TASK, EXPECTED OUTCOME\n등 5항목)",
             "매 호출마다 변경\n→ 캐시 미적중"],
        ],
        col_widths=[1.5, 2.5, 5, 5]
    )

    cap_t9 = doc.add_paragraph()
    cap_t9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t9.add_run("Table 9. 프롬프트 조립 3단계")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "1단계의 공통 정적 영역은 모든 에이전트 호출에서 동일하므로, "
        "런타임의 prefix 캐시에 가장 높은 적중률을 보인다. "
        "2단계의 에이전트별 정적 영역은 동일 에이전트를 반복 호출할 때 캐시가 적중한다. "
        "3단계의 동적 영역만이 매 호출마다 변경되어 캐시 미적중이 발생한다. "
        "이러한 설계로 LLM API 호출의 토큰 비용을 최소화할 수 있다."
    )

    add_heading(doc, "5.6 핵심 규칙 체계", level=2)

    add_body(doc,
        "DMAP은 표준 준수를 위한 명시적 규칙 체계를 정의한다."
    )

    add_body_bold(doc, "MUST 규칙 (필수)")

    make_table(doc,
        headers=["#", "규칙", "근거"],
        rows=[
            ["1", "모든 플러그인은 plugin.json과\nmarketplace.json 포함",
             "런타임 인식 진입점 및\n배포 메타데이터"],
            ["2", "모든 에이전트는 AGENT.md +\nagentcard.yaml 쌍으로 구성",
             "프롬프트/메타데이터 분리\n(경계 원칙)"],
            ["3", "tier는 HEAVY/HIGH/MEDIUM/LOW\n중 하나만 사용",
             "4-Tier 런타임 매핑 표준"],
            ["4", "위임형 스킬은 라우팅+오케스트레이션만.\n직결형 스킬은 Gateway 직접 사용 허용",
             "관심사 분리 + YAGNI"],
            ["5", "추상 선언과 구체 매핑 분리", "Dependency Inversion"],
            ["6", "setup 스킬과 core 스킬 반드시 포함.\n플러그인당 core 스킬 1개",
             "설치/라우팅 등록"],
            ["7", "AGENT.md에 도구 명세 금지.\n{tool:name} 추상 참조만 허용",
             "프롬프트/도구 분리"],
            ["8", "프롬프트 구성 순서:\n공통 정적 → 에이전트별 정적 → 동적",
             "prefix 캐시 최적화"],
        ],
        col_widths=[0.8, 6.5, 4.5]
    )

    add_body(doc, "")  # spacer

    add_body_bold(doc, "MUST NOT 규칙 (금지)")

    make_table(doc,
        headers=["#", "금지 사항", "이유"],
        rows=[
            ["1", "스킬이 직접 애플리케이션 코드 작성·수정\n(직결형 스킬의 설정 파일·문서 작업은 예외)",
             "에이전트의 역할 침범"],
            ["2", "에이전트가 직접 라우팅·오케스트레이션",
             "스킬의 역할 침범"],
            ["3", "AGENT.md에 모델명·도구명 하드코딩",
             "런타임 중립성 위반"],
            ["4", "agentcard.yaml에 프롬프트 내용 포함",
             "기계 판독용 데이터와\n프롬프트 혼재"],
            ["5", "일반 플러그인에서 Hook 사용",
             "오케스트레이션 플러그인 전용"],
            ["6", "AGENT.md와 agentcard.yaml에\n동일 정보 중복 기술",
             "경계 원칙 위반"],
        ],
        col_widths=[0.8, 7, 4]
    )

    # ===================================================================
    # SECTION 6: 사례 연구
    # ===================================================================
    add_heading(doc, "6. 사례 연구 (Case Study)", level=1)

    add_body(doc,
        "본 장에서는 DMAP 표준을 실제 적용한 두 가지 플러그인 사례를 분석한다. "
        "OMC(Oh-My-ClaudeCode) 플러그인은 오케스트레이션 전문 플러그인으로서 "
        "표준의 원천 사례이며, Abra 플러그인은 비즈니스 도메인에 표준을 적용한 사례이다."
    )

    add_heading(doc, "6.1 OMC (Oh-My-ClaudeCode) 플러그인", level=2)

    add_body(doc,
        "OMC는 Claude Code 런타임 위에서 동작하는 오케스트레이션 전문 플러그인으로, "
        "DMAP 표준의 원천(origin) 사례이다. 오케스트레이션 플러그인이므로 "
        "Hook 사용이 허용되며, 이를 통해 사용자 프롬프트 제출(UserPromptSubmit), "
        "세션 시작(SessionStart) 등 런타임 이벤트를 가로채어 "
        "자동 라우팅과 컨텍스트 주입을 수행한다."
    )

    make_table(doc,
        headers=["항목", "수치 / 내용"],
        rows=[
            ["Skills", "39개 (Core, Setup, Planning, Orchestrator, Utility 5유형)"],
            ["Agents", "35개 (12개 도메인 × 4-Tier: HEAVY/HIGH/MEDIUM/LOW)"],
            ["Hook 이벤트", "8종 (UserPromptSubmit, SessionStart 등)"],
            ["MCP Tools", "15개 (LSP 12 + AST 2 + REPL 1)"],
            ["Gateway 매핑", "tier_mapping + tool_mapping(lsp/mcp/custom) + action_mapping"],
        ],
        col_widths=[3, 11]
    )

    cap_t10 = doc.add_paragraph()
    cap_t10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t10.add_run("Table 10. OMC 플러그인 정량 데이터")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "OMC의 35개 에이전트는 분석(Architect), 실행(Executor), 탐색(Explorer), "
        "연구(Researcher), 프론트엔드(Designer), 문서(Writer), 시각(Vision), "
        "계획(Planner), 비평(Critic), 테스트(QA Tester), 보안(Security Reviewer), "
        "데이터 과학(Scientist) 등 12개 도메인에 걸쳐 분포한다. "
        "각 도메인의 에이전트는 LOW(Haiku), MEDIUM(Sonnet), HIGH(Opus), "
        "HEAVY(Opus + 대규모 예산) 티어로 변형되어, "
        "작업 복잡도에 따라 최적의 비용-역량 조합을 선택할 수 있다."
    )

    add_body(doc,
        "39개 스킬은 자동 파일럿(autopilot), 병렬 실행(ultrawork), "
        "지속 완수(ralph), 계획 수립(plan), TDD 가이드(tdd), "
        "코드 리뷰(code-review), 보안 리뷰(security-review) 등 "
        "소프트웨어 개발 전 생명주기를 포괄하며, "
        "각 스킬은 위임형 또는 직결형 경로를 통해 적절한 에이전트를 활용한다."
    )

    add_heading(doc, "6.2 Abra 플러그인", level=2)

    add_body(doc,
        "Abra는 AI Agent 개발 워크플로우를 자동화하는 비즈니스 도메인 플러그인으로, "
        "DMAP 표준의 범용성을 검증하는 사례이다. "
        "일반 플러그인이므로 Hook을 사용하지 않으며, "
        "스킬과 에이전트의 위임 체계만으로 워크플로우를 구현한다."
    )

    make_table(doc,
        headers=["항목", "내용"],
        rows=[
            ["도메인", "AI Agent 개발 파이프라인"],
            ["Skills", "setup, core, scenario, dsl-generate, prototype,\ndev-plan, develop, orchestrate"],
            ["Agents", "scenario-analyst, dsl-architect, agent-developer,\nplan-writer, prototype-runner"],
            ["워크플로우", "시나리오 정의 → DSL 생성 → 프로토타이핑 →\n개발계획 → 코드 개발"],
            ["Gateway", "install.yaml + runtime-mapping.yaml\n(tier_mapping + tool_mapping + action_mapping)"],
            ["네임스페이스", "abra:{skill-name} (스킬),\nabra:{agent}:{agent} (에이전트 FQN)"],
        ],
        col_widths=[3, 11]
    )

    cap_t11 = doc.add_paragraph()
    cap_t11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t11.add_run("Table 11. Abra 플러그인 구성")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "Abra의 워크플로우는 5단계로 구성된다. "
        "(1) scenario 스킬이 scenario-analyst 에이전트에게 "
        "AI 에이전트 시나리오 분석을 위임한다. "
        "(2) dsl-generate 스킬이 dsl-architect 에이전트에게 "
        "Dify 워크플로우 DSL 생성을 위임한다. "
        "(3) prototype 스킬이 prototype-runner 에이전트에게 "
        "프로토타이핑을 위임한다. "
        "(4) dev-plan 스킬이 plan-writer 에이전트에게 개발계획 수립을 위임한다. "
        "(5) develop 스킬이 agent-developer 에이전트에게 실제 코드 개발을 위임한다."
    )

    add_body(doc,
        "이 사례는 DMAP 표준이 소프트웨어 개발 도메인(OMC)뿐 아니라 "
        "비즈니스 워크플로우 도메인(Abra)에도 동일하게 적용 가능함을 실증한다. "
        "특히 Abra의 모든 에이전트와 스킬이 Markdown과 YAML만으로 정의되어 있으며, "
        "Python이나 TypeScript 코드가 단 한 줄도 포함되지 않는다는 점이 주목할 만하다."
    )

    # ===================================================================
    # SECTION 7: 평가
    # ===================================================================
    add_heading(doc, "7. 평가 (Evaluation)", level=1)

    add_heading(doc, "7.1 평가 프레임워크", level=2)

    add_body(doc,
        "본 연구의 평가는 7개의 비교 축을 기준으로 수행한다. "
        "Table 12는 평가 프레임워크를 정리한 것이다."
    )

    make_table(doc,
        headers=["평가 축", "측정 방법", "비교 대상"],
        rows=[
            ["진입 장벽", "첫 에이전트 정의까지 필요한\n단계 수, 필요 기술 스택",
             "LangChain, CrewAI, AutoGen"],
            ["이식성", "다른 런타임으로 이식 시\n변경 필요한 파일 수",
             "전체 재작성 vs\nruntime-mapping.yaml만 변경"],
            ["확장성", "새 에이전트/스킬 추가 시\n기존 코드 변경 필요 여부",
             "Open-Closed Principle 준수도"],
            ["토큰 효율", "동일 작업 수행 시\n소비 토큰 비교",
             "4-Tier 라우팅 +\n프롬프트 조립 캐시 효과"],
            ["관심사 분리도", "컴포넌트 간 의존 방향 수,\n순환 의존 유무",
             "Coupling 지표"],
            ["도메인 범용성", "적용 가능 도메인 수",
             "프레임워크별 사례 수"],
            ["선언 밀도", "에이전트 정의에 필요한\n라인 수 (코드 vs 선언)",
             "LoC 비교"],
        ],
        col_widths=[2.5, 5, 5.5]
    )

    cap_t12 = doc.add_paragraph()
    cap_t12.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t12.add_run("Table 12. 평가 프레임워크")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_heading(doc, "7.2 정량 비교", level=2)

    add_body(doc,
        "Table 13은 주요 프레임워크와 DMAP의 정량적 비교를 보여준다."
    )

    make_table(doc,
        headers=["평가 축", "LangChain", "CrewAI", "AutoGen", "MetaGPT", "DMAP"],
        rows=[
            ["진입 장벽\n(필요 기술)", "Python,\nSDK 학습", "Python,\nSDK 학습", "Python,\nSDK 학습",
             "Python,\nSDK 학습", "Markdown,\nYAML"],
            ["이식성\n(런타임 변경 시\n변경 파일)", "전체\n재작성", "전체\n재작성", "전체\n재작성",
             "전체\n재작성", "runtime-\nmapping.yaml\n1파일"],
            ["확장성\n(에이전트 추가 시\n기존 코드 변경)", "필요", "필요", "필요", "필요", "불필요\n(OCP 준수)"],
            ["티어 관리", "없음", "없음", "없음", "없음", "4-Tier\n자동 라우팅"],
            ["관심사 분리", "낮음", "낮음", "낮음", "중간", "높음\n(5-Layer)"],
            ["도메인 범용성", "중간", "중간", "중간", "코드\n특화", "완전 범용"],
            ["비개발자\n접근성", "불가", "불가", "불가", "불가", "가능"],
        ],
        col_widths=[2.5, 2.2, 2.0, 2.0, 2.0, 2.5]
    )

    cap_t13 = doc.add_paragraph()
    cap_t13.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t13.add_run("Table 13. 프레임워크 정량 비교")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_heading(doc, "7.3 선언 밀도 비교", level=2)

    add_body(doc,
        "에이전트 하나를 정의하는 데 필요한 라인 수를 비교하면 다음과 같다."
    )

    make_table(doc,
        headers=["프레임워크", "에이전트 정의\n방식", "필요 라인 수\n(대략)", "필요 파일 수", "비고"],
        rows=[
            ["LangChain", "Python 클래스 +\n도구 정의 + 체인", "100~300 LoC", "2~5개\n.py 파일",
             "SDK 의존성 포함"],
            ["CrewAI", "Python Agent +\nTask + Crew 정의", "50~150 LoC", "1~3개\n.py 파일",
             "데코레이터 사용"],
            ["AutoGen", "Python Agent +\n대화 설정", "80~200 LoC", "2~4개\n.py 파일",
             "대화 프로토콜 설정"],
            ["MetaGPT", "Python Role +\nAction 정의", "100~250 LoC", "3~6개\n.py 파일",
             "SOP 설정 포함"],
            ["DMAP", "Markdown +\nYAML 선언", "50~100 Lines\n(선언)", "3개\n(.md + .yaml ×2)",
             "코드 0줄,\n순수 선언"],
        ],
        col_widths=[2, 3, 2.5, 2.5, 3]
    )

    cap_t14 = doc.add_paragraph()
    cap_t14.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_t14.add_run("Table 14. 에이전트 정의 선언 밀도 비교")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run.italic = True

    add_body(doc,
        "DMAP의 에이전트 정의는 코드(code) 라인이 0줄이며, "
        "순수 선언(declaration) 라인만으로 구성된다. "
        "이는 기존 프레임워크 대비 진입 장벽이 현저히 낮으며, "
        "도메인 전문가가 개발자의 도움 없이도 에이전트를 정의할 수 있음을 의미한다. "
        "특히 Markdown은 기술 문서 작성의 사실상 표준이므로, "
        "별도의 학습 비용 없이 에이전트 프롬프트를 작성할 수 있다."
    )

    # ===================================================================
    # SECTION 8: 논의
    # ===================================================================
    add_heading(doc, "8. 논의 (Discussion)", level=1)

    add_heading(doc, "8.1 한계점", level=2)

    add_body(doc,
        "본 연구의 DMAP 아키텍처는 다음과 같은 한계점을 가진다."
    )

    add_body(doc,
        "첫째, LLM의 프롬프트 해석 정확도에 의존한다. "
        "DMAP은 에이전트의 워크플로우와 판단 기준을 자연어 프롬프트(AGENT.md)로 기술하므로, "
        "LLM이 프롬프트를 정확히 해석하고 따르는 능력에 의존한다. "
        "프롬프트의 모호성이나 LLM의 지시 따르기(instruction following) 한계는 "
        "에이전트의 행동 정확도에 직접적 영향을 미친다."
    )

    add_body(doc,
        "둘째, 복잡한 로직 표현에 한계가 있다. "
        "선언형 명세만으로는 조건부 분기, 반복, 예외 처리 등 "
        "복잡한 제어 흐름을 정밀하게 표현하기 어렵다. "
        "DMAP은 이를 프롬프트의 자연어 기술로 대체하지만, "
        "코드 수준의 정밀도를 보장하기는 어렵다."
    )

    add_body(doc,
        "셋째, 실증 범위의 제한이 있다. "
        "본 연구의 실증은 OMC(오케스트레이션)와 Abra(비즈니스 워크플로우) "
        "2개의 플러그인에 한정된다. "
        "의료, 법률, 금융 등 고위험(high-stakes) 도메인에서의 "
        "검증은 아직 이루어지지 않았다."
    )

    add_heading(doc, "8.2 위협 요소", level=2)

    add_body(doc,
        "내적 타당성 측면에서, DMAP 표준을 설계한 연구자가 "
        "동시에 실증 플러그인을 개발했다는 점은 "
        "평가의 객관성에 위협 요소로 작용할 수 있다. "
        "외적 타당성 측면에서, Claude Code라는 단일 런타임에서만 "
        "실제 실행이 검증되었으며, "
        "Codex CLI나 Gemini CLI 등 다른 런타임에서의 검증은 "
        "아키텍처적 분석에 의존한다."
    )

    add_body(doc,
        "또한 선언형 접근의 근본적 한계로, LLM 기술의 발전 속도가 "
        "선언형 표준의 갱신 속도를 초과할 경우, "
        "표준이 LLM의 최신 역량을 충분히 활용하지 못하는 "
        "지체(lag) 현상이 발생할 수 있다."
    )

    add_heading(doc, "8.3 향후 연구 방향", level=2)

    add_body(doc, "본 연구를 기반으로 다음과 같은 후속 연구를 계획한다.")

    add_numbered(doc, 1, " 현재 runtime-mapping.yaml을 수동으로 작성해야 하나, "
        "플러그인의 추상 선언과 런타임의 도구 목록을 자동으로 매칭하는 "
        "자동 매핑 엔진을 개발할 수 있다.",
        bold_prefix="런타임 자동 매핑.")
    add_numbered(doc, 2, " DMAP 표준을 커뮤니티 주도로 발전시키기 위한 "
        "개방형 표준화 프로세스와 거버넌스 구조를 제안한다.",
        bold_prefix="표준화 기구 제안.")
    add_numbered(doc, 3, " 에이전트 시스템의 성능을 객관적으로 측정할 수 있는 "
        "벤치마크 태스크 셋을 구축하여, "
        "선언형과 명령형 접근의 정량적 비교를 수행한다.",
        bold_prefix="벤치마크 구축.")
    add_numbered(doc, 4, " Claude Code 외에 Codex CLI, Gemini CLI 등 "
        "다른 런타임에서의 실제 실행 검증을 수행하여, "
        "이식성 주장을 실증적으로 강화한다.",
        bold_prefix="다중 런타임 검증.")

    # ===================================================================
    # SECTION 9: 결론
    # ===================================================================
    add_heading(doc, "9. 결론 (Conclusion)", level=1)

    add_body(doc,
        "본 연구는 LLM 기반 멀티에이전트 시스템에 소프트웨어 공학의 검증된 원칙인 "
        "Clean Architecture를 코드가 아닌 선언형 명세(Markdown + YAML)로 적용하는 "
        "DMAP(Declarative Multi-Agent Plugin) 아키텍처를 제안하였다."
    )

    add_body(doc,
        "본 연구의 핵심 기여는 네 가지로 요약된다. "
        "첫째, 마크다운과 YAML만으로 멀티에이전트 시스템을 정의하는 "
        "선언형 에이전트 아키텍처 표준을 제안하였다. "
        "둘째, Loosely Coupling, High Cohesion, Dependency Inversion 원칙을 "
        "LLM 에이전트 오케스트레이션에 체계적으로 이식하였다. "
        "셋째, 4-Tier 모델과 Gateway 매핑을 통해 "
        "런타임 중립적 추상 계층을 설계하여 이식성을 확보하였다. "
        "넷째, OMC 플러그인(39개 스킬, 35개 에이전트)과 "
        "Abra 플러그인(비즈니스 도메인)의 실제 운용을 통해 "
        "표준의 실효성을 실증하였다."
    )

    add_body(doc,
        "DMAP은 \"코드로 구현\"에서 \"선언으로 명세\"로의 패러다임 전환을 제안하며, "
        "이를 통해 (1) 비개발자의 에이전트 시스템 참여, "
        "(2) 런타임 간 이식성 확보, "
        "(3) 소프트웨어 공학 원칙에 기반한 체계적 아키텍처, "
        "(4) 도메인에 무관한 범용 적용 가능성을 실현하였다. "
        "본 연구가 LLM 에이전트 시스템의 설계와 운용에 있어 "
        "새로운 방향성을 제시할 수 있기를 기대한다."
    )

    # ===================================================================
    # REFERENCES
    # ===================================================================
    add_heading(doc, "참고문헌 (References)", level=1)

    refs = [
        '[1] Chase, H. (2022). "LangChain: Building applications with LLMs through composability." '
        'GitHub repository, https://github.com/langchain-ai/langchain.',

        '[2] Moura, J. (2023). "CrewAI: Framework for orchestrating role-playing, autonomous AI agents." '
        'GitHub repository, https://github.com/joaomdmoura/crewAI.',

        '[3] Wu, Q., Bansal, G., Zhang, J., Wu, Y., Li, B., Zhu, E., ... & Wang, C. (2023). '
        '"AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation." '
        'arXiv preprint arXiv:2308.08155.',

        '[4] Hong, S., Zhuge, M., Chen, J., Zheng, X., Cheng, Y., Zhang, C., ... & Wu, Y. (2023). '
        '"MetaGPT: Meta Programming for A Multi-Agent Collaborative Framework." '
        'arXiv preprint arXiv:2308.00352.',

        '[5] Qian, C., Cong, X., Yang, C., Chen, W., Su, Y., Xu, J., ... & Sun, M. (2023). '
        '"Communicative Agents for Software Development." '
        'arXiv preprint arXiv:2307.07924.',

        '[6] Martin, R. C. (2017). "Clean Architecture: A Craftsman\'s Guide to Software Structure and Design." '
        'Prentice Hall.',

        '[7] Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). '
        '"Design Patterns: Elements of Reusable Object-Oriented Software." '
        'Addison-Wesley.',

        '[8] Martin, R. C. (2003). "Agile Software Development: Principles, Patterns, and Practices." '
        'Prentice Hall.',

        '[9] Fowler, M. (2002). "Patterns of Enterprise Application Architecture." '
        'Addison-Wesley.',

        '[10] Beck, K. (2000). "Extreme Programming Explained: Embrace Change." '
        'Addison-Wesley.',

        '[11] Wang, L., Ma, C., Feng, X., Zhang, Z., Yang, H., Zhang, J., ... & Qin, B. (2024). '
        '"A Survey on Large Language Model based Autonomous Agents." '
        'Frontiers of Computer Science, 18(6), 186345.',

        '[12] Xi, Z., Chen, W., Guo, X., He, W., Ding, Y., Hong, B., ... & Gui, T. (2023). '
        '"The Rise and Potential of Large Language Model Based Agents: A Survey." '
        'arXiv preprint arXiv:2309.07864.',

        '[13] Talebirad, Y., & Nadiri, A. (2023). '
        '"Multi-Agent Collaboration: Harnessing the Power of Intelligent LLM Agents." '
        'arXiv preprint arXiv:2306.03314.',

        '[14] Park, J. S., O\'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P., & Bernstein, M. S. (2023). '
        '"Generative Agents: Interactive Simulacra of Human Behavior." '
        'UIST 2023.',

        '[15] Anthropic. (2024). "Claude Code: AI-powered development environment." '
        'https://docs.anthropic.com/en/docs/claude-code.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        p.paragraph_format.line_spacing = Pt(13)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)

    # ===================================================================
    # Save
    # ===================================================================
    save_path = OUTPUT_PATH
    try:
        doc.save(save_path)
    except PermissionError:
        # File is open in another program; save with alternate name
        from datetime import datetime
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = os.path.join(SCRIPT_DIR, f"DMAP-paper_{ts}.docx")
        doc.save(save_path)
        print(f"[주의] 원본 파일이 사용 중이어서 대체 파일에 저장합니다.")
    print(f"논문이 성공적으로 생성되었습니다: {save_path}")
    print(f"파일 크기: {os.path.getsize(save_path) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
