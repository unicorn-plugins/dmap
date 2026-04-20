"""
docx-build-sample.py
--------------------
DMAP 표준 DOCX 빌드 샘플 (docx-build-guide.md 4절 검증 규칙 9항 준수)

실행:    python docx-build-sample.py
산출물:  ./sample.docx (한글 폰트·헤딩·표·강조박스 포함, 2페이지 분량)
사전:    pip install python-docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────
# 가이드 4-4. 한글 폰트 헬퍼
# ─────────────────────────────────────────────────────────
def set_korean_font(run, kor="맑은 고딕", eng="Calibri", size_pt=11, bold=False, color_hex=None):
    run.font.name = eng
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), kor)
    rFonts.set(qn("w:ascii"), eng)
    rFonts.set(qn("w:hAnsi"), eng)


def set_default_korean(doc, kor="맑은 고딕"):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), kor)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")


# ─────────────────────────────────────────────────────────
# 가이드 4-7. 셀 배경색 헬퍼
# ─────────────────────────────────────────────────────────
def apply_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_korean_font(run, size_pt=size, bold=bold, color_hex=color)
    return p


# ─────────────────────────────────────────────────────────
# 가이드 4-3. 이미지 임베딩 (경로 검증)
# ─────────────────────────────────────────────────────────
def add_image_safe(doc, path_str, width_inches=6.0):
    p = Path(path_str)
    if not p.exists():
        # 샘플에서는 이미지 미존재 시 placeholder 단락으로 대체
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[이미지 placeholder: {p.name} 미존재]")
        set_korean_font(run, size_pt=10, color_hex="999999")
        return
    if p.stat().st_size == 0:
        raise ValueError(f"이미지 0바이트: {p.absolute()}")
    doc.add_picture(str(p), width=Inches(width_inches))


# ─────────────────────────────────────────────────────────
# 본문 작성
# ─────────────────────────────────────────────────────────
def build_content(doc):
    # 가이드 4-1. 헤딩은 add_heading 사용 (목차 추출 가능)
    title = doc.add_heading("DMAP Office 샘플 보고서", level=0)
    for run in title.runs:
        set_korean_font(run, size_pt=22, bold=True, color_hex="1F4E79")

    add_para(doc, "본 문서는 docx-build-guide.md 4절 검증 규칙 9항을 모두 준수하는 샘플입니다.",
             size=11, color="595959")

    # 섹션 1
    h1 = doc.add_heading("1. 개요", level=1)
    for run in h1.runs:
        set_korean_font(run, size_pt=16, bold=True, color_hex="1F4E79")

    add_para(doc, "DMAP의 MS-Office 통합 패턴은 다음 두 가지로 구성됩니다:")
    bullet_items = [
        "PPTX: spec-writer 에이전트 + 빌더 스킬 2단계 (시각 명세 분리)",
        "XLSX/DOCX: 빌더 스킬 단독 1단계 (입력 데이터/본문 직수신)",
    ]
    for item in bullet_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_korean_font(run, size_pt=11)

    # 섹션 2
    h2 = doc.add_heading("2. 핵심 통계", level=1)
    for run in h2.runs:
        set_korean_font(run, size_pt=16, bold=True, color_hex="1F4E79")

    add_para(doc, "분기별 운영 지표는 다음과 같습니다.")

    # 가이드 4-2. 표는 add_table 사용
    table_data = [
        ["분기", "매출", "비용", "이익"],
        ["1Q", "100", "60", "40"],
        ["2Q", "120", "70", "50"],
        ["3Q", "150", "80", "70"],
        ["4Q", "180", "95", "85"],
    ]
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    table.allow_autofit = False
    col_widths = [Cm(2.5), Cm(3.0), Cm(3.0), Cm(3.0)]
    for r, row_data in enumerate(table_data):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = ""  # 기존 단락 비우기
            cell.width = col_widths[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r == 0 or c == 0) else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            if r == 0:
                set_korean_font(run, size_pt=11, bold=True, color_hex="FFFFFF")
                apply_cell_shading(cell, "1F4E79")
            else:
                set_korean_font(run, size_pt=10)

    # 강조 박스 (단락 + 좌측 보더 흉내)
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(0.5)
    quote_run = quote.add_run("핵심 인사이트: 4분기 합계 이익은 245로 전년 대비 32% 성장.")
    set_korean_font(quote_run, size_pt=12, bold=True, color_hex="2E75B6")

    # 가이드 4-6. 페이지 나눔
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    # 섹션 3 (페이지 2)
    h3 = doc.add_heading("3. 빌드 검증 체크리스트", level=1)
    for run in h3.runs:
        set_korean_font(run, size_pt=16, bold=True, color_hex="1F4E79")

    checklist = [
        "헤딩이 add_heading API로 작성됨 (목차 추출 가능)",
        "한글 텍스트에 한글 폰트(맑은 고딕) 명시됨",
        "표는 add_table API 사용, 헤더 행 별도 스타일 적용",
        "셀 배경색이 헤더에 적용됨",
        "이미지 경로 존재 검증 후 임베딩 (또는 placeholder)",
        "페이지 나눔이 의도한 위치에 적용됨",
        "본문 단락 간격·줄 간격 일관됨",
        "빌드 스크립트 종료 코드 0",
        "Word/LibreOffice에서 시각적 검증 통과",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_korean_font(run, size_pt=11)

    # 이미지 placeholder 예시
    h4 = doc.add_heading("4. 다이어그램 (예시)", level=1)
    for run in h4.runs:
        set_korean_font(run, size_pt=16, bold=True, color_hex="1F4E79")

    add_image_safe(doc, "images/architecture.png", width_inches=6.0)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run("그림 1. DMAP Office 통합 아키텍처")
    set_korean_font(cap_run, size_pt=9, color_hex="595959")


# ─────────────────────────────────────────────────────────
# 가이드 4-8. 빌드 스크립트 진입점
# ─────────────────────────────────────────────────────────
def build():
    doc = Document()
    set_default_korean(doc)
    build_content(doc)

    output = Path(__file__).parent / "sample.docx"
    doc.save(str(output))

    # 가이드 4-9. 자가 검증 (최소 항목)
    if not output.exists() or output.stat().st_size == 0:
        raise RuntimeError(f"산출물 검증 실패: {output}")

    print(f"[OK] saved: {output} ({output.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(build())
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
